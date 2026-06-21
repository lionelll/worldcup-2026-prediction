#!/usr/bin/env python3
"""Keep the user-edited front matter and replace Chapter 3 onward."""

from __future__ import annotations

import argparse
import copy
import os
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PR = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W, "r": R}


def text_of(element):
    try:
        values = element.xpath(".//w:t/text()", namespaces=NS)
    except TypeError:
        values = element.xpath(".//w:t/text()")
    return "".join(values)


def chapter_three_index(body):
    candidates = []
    for index, child in enumerate(body):
        if child.tag == f"{{{W}}}p" and "第三章" in text_of(child):
            candidates.append(index)
    if candidates:
        return candidates[-1]
    raise RuntimeError("未找到第三章边界")


def truncate_at_marker(element, marker="第三章"):
    full = text_of(element)
    keep = full.find(marker)
    if keep < 0:
        return
    remaining = keep
    for node in element.xpath(".//w:t", namespaces=NS):
        value = node.text or ""
        if remaining >= len(value):
            remaining -= len(value)
        elif remaining > 0:
            node.text = value[:remaining]
            remaining = 0
        else:
            node.text = ""


def merge_packages(front_path: Path, generated_path: Path, output_path: Path):
    with tempfile.TemporaryDirectory() as temp:
        root = Path(temp)
        front_dir, generated_dir = root / "front", root / "generated"
        with zipfile.ZipFile(front_path) as archive:
            archive.extractall(front_dir)
        with zipfile.ZipFile(generated_path) as archive:
            archive.extractall(generated_dir)

        parser = etree.XMLParser(remove_blank_text=False)
        front_xml = etree.parse(str(front_dir / "word/document.xml"), parser)
        generated_xml = etree.parse(str(generated_dir / "word/document.xml"), parser)
        front_body = front_xml.find(f".//{{{W}}}body")
        generated_body = generated_xml.find(f".//{{{W}}}body")
        front_cut = chapter_three_index(front_body)
        generated_cut = chapter_three_index(generated_body)

        front_elements = [copy.deepcopy(x) for x in list(front_body)[: front_cut + 1]]
        truncate_at_marker(front_elements[-1])
        if not text_of(front_elements[-1]).strip():
            front_elements.pop()

        front_rels_path = front_dir / "word/_rels/document.xml.rels"
        generated_rels_path = generated_dir / "word/_rels/document.xml.rels"
        front_rels = etree.parse(str(front_rels_path), parser)
        generated_rels = etree.parse(str(generated_rels_path), parser)
        source_rels = {x.get("Id"): x for x in front_rels.getroot()}
        used_ids = {x.get("Id") for x in generated_rels.getroot()}
        next_id = 1

        def allocate_id():
            nonlocal next_id
            while f"rId{next_id}" in used_ids:
                next_id += 1
            value = f"rId{next_id}"
            used_ids.add(value)
            next_id += 1
            return value

        id_map = {}
        media_counter = 1
        for element in front_elements:
            for node in element.iter():
                for attr in (f"{{{R}}}embed", f"{{{R}}}id", f"{{{R}}}link"):
                    old_id = node.get(attr)
                    if not old_id:
                        continue
                    if old_id not in id_map:
                        source = source_rels.get(old_id)
                        if source is None:
                            continue
                        new_id = allocate_id()
                        relation = copy.deepcopy(source)
                        relation.set("Id", new_id)
                        target = relation.get("Target")
                        if relation.get("TargetMode") != "External" and target:
                            source_part = (front_dir / "word" / target).resolve()
                            if source_part.exists():
                                suffix = source_part.suffix
                                new_target = f"media/front_{media_counter}{suffix}"
                                media_counter += 1
                                destination = generated_dir / "word" / new_target
                                destination.parent.mkdir(parents=True, exist_ok=True)
                                shutil.copy2(source_part, destination)
                                relation.set("Target", new_target)
                        generated_rels.getroot().append(relation)
                        id_map[old_id] = new_id
                    node.set(attr, id_map[old_id])

        tail = [copy.deepcopy(x) for x in list(generated_body)[generated_cut:]]
        for child in list(generated_body):
            generated_body.remove(child)
        for child in front_elements + tail:
            generated_body.append(child)

        generated_xml.write(str(generated_dir / "word/document.xml"), xml_declaration=True, encoding="UTF-8", standalone="yes")
        generated_rels.write(str(generated_rels_path), xml_declaration=True, encoding="UTF-8", standalone="yes")
        shutil.copy2(front_dir / "word/styles.xml", generated_dir / "word/styles.xml")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in generated_dir.rglob("*"):
                if path.is_file():
                    archive.write(path, path.relative_to(generated_dir))


def set_east_asia(run, font_name):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def remove_numbering(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    fill = tc_pr.find(qn("w:shd"))
    if fill is None:
        fill = OxmlElement("w:shd")
        tc_pr.append(fill)
    fill.set(qn("w:fill"), color)


def replace_across_runs(paragraph, old, new):
    full = "".join(run.text for run in paragraph.runs)
    if old not in full:
        return False
    updated = full.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def format_document(path: Path):
    doc = Document(path)
    in_tail = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "第三章" in text:
            in_tail = True
        if not in_tail:
            if text.startswith("本文基于历史国际足球比赛数据"):
                replace_across_runs(
                    paragraph,
                    "本文基于历史国际足球比赛数据",
                    "本文基于当前阶段性快照和历史国际足球比赛数据",
                )
            replace_across_runs(paragraph, "准确率为 55.3%", "准确率为 55.31%")
            replace_across_runs(paragraph, "已录入的 30 场", "已录入的 32 场")
            replace_across_runs(paragraph, "方向命中率为 56.7%", "方向命中率为 56.25%")
            replace_across_runs(
                paragraph,
                "截至 6 月 20 日 09:18 的动态预测",
                "截至 6 月 20 日共 32 场完赛后的动态预测",
            )
            continue

        if text.startswith("第") and "章" in text[:12]:
            paragraph.style = doc.styles["Normal"]
            remove_numbering(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(12)
            for run in paragraph.runs:
                set_east_asia(run, "Heiti SC Light")
                run.font.size = Pt(16)
                run.font.bold = True
            continue
        if text and len(text.split(" ", 1)[0].split(".", 2)) >= 2 and text[0].isdigit() and "." in text[:6]:
            paragraph.style = doc.styles["Normal"]
            remove_numbering(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                set_east_asia(run, "宋体")
                run.font.size = Pt(14)
                run.font.bold = True
            continue
        if not text:
            continue
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(5)
            is_caption = text.startswith(("表 ", "图 ", "表", "图"))
            for run in paragraph.runs:
                if is_caption:
                    set_east_asia(run, "宋体")
                    run.font.size = Pt(10.5)
                    run.font.bold = True
            continue
        paragraph.style = doc.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)
        for run in paragraph.runs:
            set_east_asia(run, "宋体")
            run.font.size = Pt(12)

    passed = False
    for table in doc.tables:
        prior_text = ""
        xml = table._tbl
        for sibling in xml.itersiblings(preceding=True):
            if sibling.tag == qn("w:p"):
                prior_text = text_of(sibling)
                if "第三章" in prior_text:
                    passed = True
                    break
        if not passed:
            continue
        for cell in table.rows[0].cells:
            shade(cell, "D9EAF7")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    set_east_asia(run, "宋体")
                    run.font.color.rgb = RGBColor(31, 78, 121)
        for row in table.rows[1:]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_east_asia(run, "宋体")

    toc_rows = [
        ("第一章　绪论", "5"),
        ("第二章　数据来源与预处理", "8"),
        ("第三章　探索性数据分析", "12"),
        ("第四章　研究方法与模型实现", "16"),
        ("第五章　模型训练、验证与诊断", "18"),
        ("第六章　截至 2026-06-20 的赛中检验", "25"),
        ("第七章　世界杯模拟结果", "28"),
        ("第八章　赛中更新的影响", "42"),
        ("第九章　结论、局限与改进方案", "44"),
        ("参考文献与数据来源", "46"),
        ("附录 A　R 代码与复现说明", "48"),
        ("附录 B　运行一致性检查", "48"),
        ("附录 C　网页展示与打开方式", "49"),
    ]
    if doc.tables and len(doc.tables[0].rows) == len(toc_rows):
        for row, (label, page) in zip(doc.tables[0].rows, toc_rows):
            for cell, value in zip(row.cells, (label, page)):
                paragraph = cell.paragraphs[0]
                if paragraph.runs:
                    paragraph.runs[0].text = value
                    for run in paragraph.runs[1:]:
                        run.text = ""
                else:
                    paragraph.add_run(value)

    doc.save(path)


def restore_image_relationships(formatted_path: Path, reference_path: Path):
    """python-docx collapses copied image rels on save; restore them from the merged package."""
    with tempfile.TemporaryDirectory() as temp:
        root = Path(temp)
        formatted_dir, reference_dir = root / "formatted", root / "reference"
        with zipfile.ZipFile(formatted_path) as archive:
            archive.extractall(formatted_dir)
        with zipfile.ZipFile(reference_path) as archive:
            archive.extractall(reference_dir)

        parser = etree.XMLParser(remove_blank_text=False)
        formatted_xml = etree.parse(str(formatted_dir / "word/document.xml"), parser)
        reference_xml = etree.parse(str(reference_dir / "word/document.xml"), parser)
        formatted_images = formatted_xml.xpath("//*[@r:embed]", namespaces=NS)
        reference_images = reference_xml.xpath("//*[@r:embed]", namespaces=NS)
        if len(formatted_images) != len(reference_images):
            raise RuntimeError(
                f"图片节点数量不一致：格式化文档 {len(formatted_images)}，参考包 {len(reference_images)}"
            )
        for formatted, reference in zip(formatted_images, reference_images):
            formatted.set(f"{{{R}}}embed", reference.get(f"{{{R}}}embed"))
        formatted_xml.write(
            str(formatted_dir / "word/document.xml"),
            xml_declaration=True,
            encoding="UTF-8",
            standalone="yes",
        )

        reference_rels_path = reference_dir / "word/_rels/document.xml.rels"
        reference_rels = etree.parse(str(reference_rels_path), parser)
        referenced_ids = {
            node.get(f"{{{R}}}embed") for node in formatted_images if node.get(f"{{{R}}}embed")
        }
        for relation in list(reference_rels.getroot()):
            if relation.get("Target") == "NULL" and relation.get("Id") not in referenced_ids:
                reference_rels.getroot().remove(relation)
        reference_rels.write(
            str(formatted_dir / "word/_rels/document.xml.rels"),
            xml_declaration=True,
            encoding="UTF-8",
            standalone="yes",
        )

        media_dir = formatted_dir / "word/media"
        shutil.rmtree(media_dir, ignore_errors=True)
        shutil.copytree(reference_dir / "word/media", media_dir)

        rebuilt = formatted_path.with_suffix(".rebuilt.docx")
        with zipfile.ZipFile(rebuilt, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in formatted_dir.rglob("*"):
                if path.is_file():
                    archive.write(path, path.relative_to(formatted_dir))
        os.replace(rebuilt, formatted_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--front", type=Path, required=True)
    parser.add_argument("--generated", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    interim = args.output.with_suffix(".merged.docx")
    merge_packages(args.front, args.generated, interim)
    shutil.copy2(interim, args.output)
    format_document(args.output)
    restore_image_relationships(args.output, interim)
    interim.unlink(missing_ok=True)
    print(args.output)


if __name__ == "__main__":
    main()
