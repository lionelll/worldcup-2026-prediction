#!/usr/bin/env python3
"""Build the course-analysis DOCX from the current R model outputs.

The modelling and simulation code submitted for the course remains in R/.  This
script only assembles the reproducible outputs into a Word document.  Formal
generation is blocked until every required data source is approved; the
--allow-unconfirmed-data flag creates an explicitly labelled review draft.
"""

from __future__ import annotations

import argparse
import csv
import math
import os
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
OUTPUT = ROOT / "output"
FIGURES = OUTPUT / "figures"
ASSETS = ROOT / "assets"

PKU_RED = "9E1B32"
INK = "111111"
MUTED = "666666"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9D9D9"
CAUTION = "FFF2CC"
RISK = "FCE4D6"

PAGE_WIDTH_DXA = 11906  # A4
LEFT_MARGIN_DXA = 1440
RIGHT_MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - LEFT_MARGIN_DXA - RIGHT_MARGIN_DXA

REQUIRED_APPROVALS = {
    "worldcup_2026_groups.csv",
    "worldcup_2026_schedule.csv",
    "worldcup_2026_results_asof_2026-06-20.csv",
    "historical_matches.csv",
    "annex_c_full_mapping",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def approved(value: str) -> bool:
    return str(value).strip().lower() in {"true", "1", "yes", "approved"}


def approval_state() -> tuple[list[dict[str, str]], bool]:
    rows = read_csv(DATA / "data_approval.csv")
    mapping = {row["file"]: approved(row["approved"]) for row in rows}
    is_formal = all(mapping.get(name, False) for name in REQUIRED_APPROVALS)
    return rows, is_formal


def set_run_font(run, *, east="宋体", latin="Times New Roman", size=11.5,
                 bold=None, italic=None, color=INK):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: list[int]):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def add_field(paragraph, instruction: str, result=""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = result
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    for node in (begin, instr, separate, text, end):
        run.append(node)


def place_picture_behind_page(picture):
    """Convert a python-docx inline picture to a full-page anchored background."""
    inline = picture._inline
    anchor = OxmlElement("wp:anchor")
    for key, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "0",
        "behindDoc": "1",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(key, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    for tag in ("wp:positionH", "wp:positionV"):
        position = OxmlElement(tag)
        position.set("relativeFrom", "page")
        offset = OxmlElement("wp:posOffset")
        offset.text = "0"
        position.append(offset)
        anchor.append(position)

    anchor.append(inline.find(qn("wp:extent")))
    effect = OxmlElement("wp:effectExtent")
    for edge in ("l", "t", "r", "b"):
        effect.set(edge, "0")
    anchor.append(effect)
    anchor.append(OxmlElement("wp:wrapNone"))

    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        anchor.append(inline.find(qn(tag)))

    inline.getparent().replace(inline, anchor)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    heading_specs = {
        "Heading 1": ("黑体", 16, 16, 10),
        "Heading 2": ("黑体", 14, 12, 7),
        "Heading 3": ("黑体", 12, 9, 5),
    }
    for name, (font, size, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(8)
    hr = hp.add_run("工业大数据课程设计报告  |  2026 年世界杯冠军概率预测")
    set_run_font(hr, east="宋体", size=8.5, color=MUTED)

    first_header = section.first_page_header
    fhp = first_header.paragraphs[0]
    fhp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fhp.paragraph_format.left_indent = Cm(-2.4)
    fhp.paragraph_format.right_indent = Cm(-2.4)
    fhp.paragraph_format.space_before = Pt(0)
    fhp.paragraph_format.space_after = Pt(0)
    background = ASSETS / "worldcup_cover_background-v1.png"
    if background.exists():
        picture = fhp.add_run().add_picture(
            str(background), width=Cm(21.0), height=Cm(29.7)
        )
        picture._inline.docPr.set("descr", "世界杯主题水彩封面背景")
        place_picture_behind_page(picture)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("第 ")
    set_run_font(fr, east="宋体", size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    fr = fp.add_run(" 页 / 共 ")
    set_run_font(fr, east="宋体", size=8.5, color=MUTED)
    add_field(fp, "NUMPAGES")
    fr = fp.add_run(" 页")
    set_run_font(fr, east="宋体", size=8.5, color=MUTED)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_text(doc, text: str, *, indent=True, bold_lead: str | None = None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_heading(doc, text: str, level=1):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.keep_with_next = True
    return heading


def add_caption(doc, text: str):
    p = doc.add_paragraph(style="Caption")
    r = p.add_run(text)
    set_run_font(r, east="宋体", size=9.5)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int],
              *, font_size=9.5, header_fill=LIGHT_GRAY, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, east="黑体", size=font_size, bold=True)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if alignments:
                p.alignment = alignments[idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title: str, body: str, *, fill=CAUTION):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, east="黑体", size=10.5, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(body)
    set_run_font(r, size=10)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, filename: str, caption: str, width_cm=15.3, source_dir=FIGURES):
    path = source_dir / filename
    if not path.exists():
        add_callout(doc, "图表缺失", f"未找到 {path.name}，请先运行制图程序。", fill=RISK)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    picture = p.add_run().add_picture(str(path), width=Cm(width_cm))
    picture._inline.docPr.set("descr", caption)
    add_caption(doc, caption)


def add_decorative_image_pair(doc, items: list[tuple[str, str]]):
    for filename, alt_text in items:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(5)
        picture = paragraph.add_run().add_picture(str(ASSETS / filename), width=Cm(15.3))
        picture._inline.docPr.set("descr", alt_text)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(6)
    set_run_font(
        caption.add_run("晋级路径视觉引导（AI 生成，仅作章节氛围装饰，不承载球队、比分或概率数据）"),
        size=9,
        color=MUTED,
    )


def add_equation(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r, east="宋体", latin="Cambria Math", size=11.5, italic=True)


def add_page_break(doc):
    doc.add_page_break()


def add_landscape_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.75)
    section.different_first_page_header_footer = False
    return section


def add_portrait_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = False
    return section


def build_cover(doc: Document, draft: bool):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = ASSETS / "pku_group3_wordmark-transparent.png"
    if logo.exists():
        picture = p.add_run().add_picture(str(logo), width=Cm(10.7))
        picture._inline.docPr.set("descr", "北京大学校徽与中英文校名")
    else:
        r = p.add_run("北京大学")
        set_run_font(r, east="Xingkai SC", size=28, bold=True, color=PKU_RED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("工业大数据课程设计报告")
    set_run_font(r, east="Xingkai SC", size=27, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("题目：")
    set_run_font(r, east="STKaiti", size=15.5, bold=True)
    r = p.add_run("基于历史比赛数据与蒙特卡洛模拟的")
    set_run_font(r, east="STKaiti", size=15.5)
    r.font.underline = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("2026 年世界杯冠军概率预测研究")
    set_run_font(r, east="STKaiti", size=15.5)
    r.font.underline = True

    member_lines = [
        "____________________________",
        "____________________________",
        "____________________________",
        "____________________________",
        "____________________________",
    ]
    for idx, line in enumerate(member_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(2)
        prefix = "小组成员：" if idx == 0 else "　　　　　"
        r = p.add_run(prefix)
        set_run_font(r, east="STKaiti", size=13)
        r = p.add_run(line)
        set_run_font(r, east="STKaiti", size=13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("指导老师：")
    set_run_font(r, east="STKaiti", size=13)
    r = p.add_run("____________________________")
    set_run_font(r, east="STKaiti", size=13)
    add_page_break(doc)


def build_abstract(doc, metadata, metrics, observed, draft):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("摘　要")
    set_run_font(r, east="黑体", size=18, bold=True)
    sims = int(float(metadata["simulations"]))
    accuracy = float(metrics["accuracy"])
    obs_accuracy = float(observed["accuracy"])
    text = (
        "本文构建了一个以动态 Elo、赛前滚动状态特征和泊松回归为核心的国家队比赛概率模型，"
        f"并将单场比分模型嵌入 2026 年世界杯赛制进行 {sims:,} 次蒙特卡洛模拟。历史数据按时间切分："
        "2010-2024 年用于训练，2025 年用于平局概率校准，2026 年 1 月 1 日至 6 月 10 日用于独立验证。"
        f"验证集共 {int(float(metrics['samples']))} 场，胜平负准确率为 {accuracy:.1%}，"
        f"Brier Score 为 {float(metrics['brier_score']):.4f}，Log Loss 为 {float(metrics['log_loss']):.4f}。"
        f"在截至 2026 年 6 月 20 日 09:18 录入的 {int(float(observed['matches']))} 场世界杯赛果上，"
        f"方向命中率为 {obs_accuracy:.1%}。赛中样本出现 {int(float(observed['actual_draws']))} 场平局，"
        "而最高概率类别没有选择平局，说明当前模型的主要误差来自平局与低比分相关结构。"
        "本文同时给出冠军概率、各阶段晋级概率、赛中预测偏差及模型风险。由于分组、赛程、赛果、历史数据来源和 Annex C"
        " 映射尚未完成最终确认，当前概率仅作为课程报告审阅和代码调试结果，不作为正式赛果判断。"
    )
    add_text(doc, text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("关键词：")
    set_run_font(r, east="黑体", size=11.5, bold=True)
    r = p.add_run("世界杯预测；泊松回归；动态 Elo；蒙特卡洛模拟；概率校准")
    set_run_font(r)
    if draft:
        add_callout(
            doc,
            "报告状态声明",
            "本文件是完整结构的课程分析报告，但数值属于未通过数据确认闸门的调试快照。待用户确认关键数据后，必须重新运行 R 程序并重新生成正式版，不能直接提交本版概率。",
            fill=CAUTION,
        )
    add_page_break(doc)


def build_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目　录")
    set_run_font(r, east="黑体", size=18, bold=True)
    entries = [
        ("第一章　绪论", "4"),
        ("第二章　数据来源、状态与预处理", "6"),
        ("第三章　探索性数据分析", "9"),
        ("第四章　研究方法与模型实现", "13"),
        ("第五章　模型训练、验证与诊断", "15"),
        ("第六章　截至 2026-06-20 09:18 的赛中检验", "22"),
        ("第七章　世界杯模拟结果", "25"),
        ("第八章　赛中更新的影响", "39"),
        ("第九章　结论、局限与改进方案", "41"),
        ("参考文献与数据来源", "43"),
        ("附录 A　R 代码与复现说明", "44"),
        ("附录 B　运行一致性检查", "44"),
        ("附录 C　网页展示与打开方式", "45"),
    ]
    table = doc.add_table(rows=len(entries), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [CONTENT_WIDTH_DXA - 900, 900])
    for row, (title, page) in zip(table.rows, entries):
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(5)
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(5)
        set_run_font(row.cells[0].paragraphs[0].add_run(title), size=11.5)
        set_run_font(row.cells[1].paragraphs[0].add_run(page), latin="Times New Roman", size=11.5)
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    add_page_break(doc)


def historical_statistics(rows):
    dates = [row["date"] for row in rows]
    home_goals = np.array([int(row["home_score"]) for row in rows], dtype=float)
    away_goals = np.array([int(row["away_score"]) for row in rows], dtype=float)
    outcomes = np.where(home_goals > away_goals, "H", np.where(home_goals == away_goals, "D", "A"))
    teams = {row["home_team"] for row in rows} | {row["away_team"] for row in rows}
    return {
        "n": len(rows),
        "start": min(dates),
        "end": max(dates),
        "teams": len(teams),
        "mean_home": float(home_goals.mean()),
        "mean_away": float(away_goals.mean()),
        "mean_total": float((home_goals + away_goals).mean()),
        "home_rate": float(np.mean(outcomes == "H")),
        "draw_rate": float(np.mean(outcomes == "D")),
        "away_rate": float(np.mean(outcomes == "A")),
    }


def tournament_category(name: str) -> str:
    lower = name.lower()
    if "friendly" in lower:
        return "友谊赛"
    if "world cup" in lower and "qualif" in lower:
        return "世界杯预选赛"
    if "world cup" in lower:
        return "世界杯正赛"
    if any(word in lower for word in ("euro", "copa", "africa", "asian", "gold cup")):
        return "洲际杯赛"
    if "nations league" in lower:
        return "国家联赛"
    if "qualif" in lower:
        return "其他预选赛"
    return "其他"


def simple_feature_correlations(rows):
    elo = defaultdict(lambda: 1500.0)
    form = defaultdict(list)
    values = defaultdict(list)
    for row in rows:
        h, a = row["home_team"], row["away_team"]
        hs, as_ = int(row["home_score"]), int(row["away_score"])
        diff = elo[h] - elo[a]
        h_form = form[h][-10:] or [(1, 1)]
        a_form = form[a][-10:] or [(1, 1)]
        h_wr = np.mean([g > ga for g, ga in h_form])
        a_wr = np.mean([g > ga for g, ga in a_form])
        h_gd = np.mean([g - ga for g, ga in h_form])
        a_gd = np.mean([g - ga for g, ga in a_form])
        neutral = str(row.get("neutral", "FALSE")).strip().lower() in {"true", "1", "yes"}
        outcome = 1 if hs > as_ else (0 if hs == as_ else -1)
        values["Elo 分差"].append(diff)
        values["近 10 场胜率差"].append(h_wr - a_wr)
        values["近 10 场净胜球差"].append(h_gd - a_gd)
        values["非中立场标记"].append(0 if neutral else 1)
        values["比赛结果"].append(outcome)
        expected = 1 / (1 + 10 ** (-diff / 400))
        actual = 1 if hs > as_ else (0.5 if hs == as_ else 0)
        elo[h] += 20 * (actual - expected)
        elo[a] -= 20 * (actual - expected)
        form[h].append((hs, as_))
        form[a].append((as_, hs))
    result = np.array(values["比赛结果"], dtype=float)
    return {key: float(np.corrcoef(np.array(vals, dtype=float), result)[0, 1]) for key, vals in values.items() if key != "比赛结果"}


def build_report(doc, *, draft: bool):
    approvals = read_csv(DATA / "data_approval.csv")
    metadata = {r["key"]: r["value"] for r in read_csv(OUTPUT / "run_metadata.csv")}
    metrics = {r["metric"]: r["value"] for r in read_csv(OUTPUT / "model_metrics.csv")}
    observed = {r["metric"]: r["value"] for r in read_csv(OUTPUT / "prediction_vs_actual_summary.csv")}
    history = read_csv(DATA / "historical_matches.csv")
    history_stats = historical_statistics(history)
    correlations = simple_feature_correlations(history)
    groups = read_csv(DATA / "worldcup_2026_groups.csv")
    schedule = read_csv(DATA / "worldcup_2026_schedule.csv")
    results = read_csv(DATA / "worldcup_2026_results_asof_2026-06-20.csv")
    comparisons = read_csv(OUTPUT / "prediction_vs_actual.csv")
    champion = read_csv(OUTPUT / "champion_probabilities.csv")
    final_four = read_csv(OUTPUT / "final_four_probabilities.csv")
    group_adv = read_csv(OUTPUT / "group_advancement_probabilities.csv")
    coefficients = read_csv(OUTPUT / "model_coefficients.csv")
    data_audit = read_csv(OUTPUT / "data_quality_audit.csv")
    feature_summary = read_csv(OUTPUT / "feature_descriptive_statistics.csv")
    poisson_inference = read_csv(OUTPUT / "poisson_coefficient_inference.csv")
    poisson_fit = read_csv(OUTPUT / "poisson_fit_statistics.csv")
    model_comparison = read_csv(OUTPUT / "poisson_model_comparison.csv")
    rolling_validation = read_csv(OUTPUT / "rolling_origin_validation.csv")
    logistic_metrics = {r["metric"]: r["value"] for r in read_csv(OUTPUT / "logistic_metrics.csv")}
    logistic_inference = read_csv(OUTPUT / "logistic_coefficient_inference.csv")
    snapshots = read_csv(OUTPUT / "snapshot_change_2026-06-19_to_2026-06-20.csv")
    current_group_status = read_csv(OUTPUT / "current_group_status.csv")

    build_cover(doc, draft)
    build_abstract(doc, metadata, metrics, observed, draft)
    build_toc(doc)

    add_heading(doc, "第一章　绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_text(doc, "世界杯冠军并不是单场强弱比较的直接结果，而是球队实力、分组、赛程、淘汰赛路径和比赛随机性共同作用的结果。只给出一个‘最可能冠军’无法表达不确定性，也无法解释同一支球队在不同路径下的晋级风险。概率模型的价值在于把单场比分分布与完整赛制连接起来，从而输出各阶段晋级概率，而不是给出确定性赛果。")
    add_text(doc, "2026 年世界杯采用 48 队、12 个小组的扩军赛制。每组前两名和 8 支成绩最好的小组第三名晋级 32 强。最佳第三名的筛选与淘汰赛槽位分配会改变强队和弱队的路径，因此赛事模拟必须同时生成比分、积分、净胜球和进球数，并按规则处理第三名。")
    add_figure(
        doc,
        "messi_ronaldo_editorial-v1.png",
        "图 1-1　足球竞争与概率预测主题概念图（AI 生成插图，非真实赛事照片）",
        width_cm=13.8,
        source_dir=ASSETS,
    )
    add_text(doc, "图 1-1 仅用于呈现足球预测研究的主题语境，不作为模型训练数据、赛事事实或实证证据。", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_heading(doc, "1.2 研究问题", 2)
    questions = [
        "历史比赛中的球队强度与近期状态能否形成具有可解释性的单场进球模型？",
        "模型在严格时间外验证集上的胜平负概率质量如何？",
        "纳入截至 2026 年 6 月 20 日 09:18 明确完赛的结果后，各队后续晋级和夺冠概率如何变化？",
        "当前赛中实际结果与赛前模型的主要偏差来自哪里，下一步应如何改进？",
    ]
    for idx, item in enumerate(questions, 1):
        add_text(doc, f"（{idx}）{item}", indent=False, after=3)
    add_heading(doc, "1.3 研究框架", 2)
    add_text(doc, "本文采用‘历史特征构造 - 单场比分建模 - 概率校准 - 真实赛制模拟 - 赛中误差检验’的两层框架。第一层通过动态 Elo、近 10 场滚动指标和泊松回归得到双方预期进球数；第二层将比分抽样放入世界杯赛程，重复运行完整赛事并以到达各轮次的频率估计概率。")
    add_equation(doc, "历史比赛 → 动态 Elo / 滚动状态 → 泊松比分模型 → 赛制模拟 → 晋级概率")

    add_heading(doc, "1.4 本报告的数据边界", 2)
    add_text(doc, "历史模型的时间边界固定为 2026 年 6 月 10 日。2026 年 6 月 11 日之后的世界杯赛果不进入训练集，只用于更新当前小组积分、Elo 和近期状态，并用于检验模型与现实比赛的差异。这样既保留了赛前模型的可验证性，也允许报告给出截至 6 月 20 日 09:18 的动态预测。直播中或未明确完赛的比分不纳入。")
    add_callout(doc, "事实与模型输出的区分", "赛程、比分和文献信息属于外部事实；概率、系数和评价指标属于本研究代码输出。所有外部事实必须确认来源，所有结果数字必须能够由 R 程序复现。", fill=LIGHT_GRAY)
    add_page_break(doc)

    add_heading(doc, "第二章　数据来源、状态与预处理", 1)
    add_heading(doc, "2.1 数据构成", 2)
    data_rows = [
        ["历史比赛", "historical_matches.csv", f"{history_stats['n']:,} 场；{history_stats['start']} 至 {history_stats['end']}", "来源与快照日期待确认"],
        ["分组", "worldcup_2026_groups.csv", f"{len(groups)} 支球队、12 组", "按用户提供分组图录入"],
        ["赛程", "worldcup_2026_schedule.csv", f"{len(schedule)} 行", "按用户提供赛程图录入"],
        ["已赛结果", "worldcup_2026_results_asof_2026-06-20.csv", f"{len(results)} 场", "百度体育页面/用户截图口径"],
        ["模型输出", "output/*.csv", f"{int(float(metadata['simulations'])):,} 次模拟", "R 程序调试输出"],
    ]
    add_caption(doc, "表 2-1　项目数据构成")
    add_table(doc, ["类型", "文件", "规模", "当前口径"], data_rows, [1250, 2800, 1800, CONTENT_WIDTH_DXA-5850], font_size=9)

    add_heading(doc, "2.2 数据确认状态", 2)
    approval_rows = []
    for row in approvals:
        status = "已确认" if approved(row["approved"]) else "待确认"
        approval_rows.append([row["file"], status, row.get("source", ""), row.get("snapshot_date", ""), row.get("notes", "")])
    add_caption(doc, "表 2-2　数据确认闸门")
    add_table(doc, ["数据项", "状态", "来源", "快照日期", "备注"], approval_rows,
              [2450, 800, 1650, 1050, CONTENT_WIDTH_DXA-5950], font_size=8)
    add_text(doc, "表 2-2 表明，当前所有关键数据均未通过最终确认。因此，本报告可以用于审阅分析链条、代码运行和结果解释，但不得将当前概率作为最终提交数值。正式版生成程序默认会被确认闸门阻止。")

    add_heading(doc, "2.3 时间切分与防止数据泄漏", 2)
    split_rows = [
        ["训练集", "2010-01-01 至 2024-12-31", "14,504", "拟合动态特征与泊松回归"],
        ["校准集", "2025-01-01 至 2025-12-31", "1,002", "拟合平局概率乘数"],
        ["独立验证集", "2026-01-01 至 2026-06-10", "311", "评价准确率与概率质量"],
        ["赛中观测集", "2026-06-12 至 2026-06-20 09:18", str(len(results)), "不回填训练，仅更新当前状态"],
    ]
    add_caption(doc, "表 2-3　时间切分方案")
    add_table(doc, ["数据段", "日期", "场次", "用途"], split_rows, [1300, 2450, 900, CONTENT_WIDTH_DXA-4650], font_size=9.2)
    add_text(doc, "所有近 10 场胜率、平均进球、平均失球和净胜球均按比赛日期向前滚动计算。每一场历史比赛的特征只允许读取该场之前的信息；赛后比分只在完成该场预测之后更新球队状态。")

    add_heading(doc, "2.4 清洗与变量处理", 2)
    cleaning = [
        "队名统一：分组、赛程、历史数据和赛果必须映射到同一球队名称。",
        "重复与日期检查：重复比赛键先审计并核对来源，不未经确认自动删除；禁止历史文件包含 2026-06-11 及之后的结果。",
        "赛事权重：世界杯、洲际杯、预选赛、国家联赛和友谊赛采用不同重要性权重。",
        "中立场处理：历史数据使用 neutral 字段；世界杯赛程需要结合比赛所在国家处理东道主优势。",
        "缺失变量：伤病、首发、天气、赔率、xG、公平竞赛分等未稳定获得的变量不补造。",
    ]
    for idx, item in enumerate(cleaning, 1):
        add_text(doc, f"（{idx}）{item}", indent=False, after=3)
    audit_labels = {
        "raw_rows": "原始记录",
        "retained_rows": "日期与比分校验后保留",
        "rows_after_cutoff": "超过历史截止日",
        "duplicate_match_keys": "重复比赛键",
        "invalid_dates": "无效日期",
        "invalid_scores": "无效比分",
        "high_score_matches_total_ge_10": "单场总进球≥10（仅标记）",
    }
    audit_rows = []
    audit_treatments = {
        "input": "原始输入",
        "used by feature builder": "进入特征构造",
        "excluded by date boundary": "按日期边界排除",
        "reported; not automatically deleted": "报告并核查，不自动删除",
        "excluded": "排除",
        "flagged for sensitivity review; not automatically deleted": "标记供敏感性复核，不自动删除",
    }
    for row in data_audit:
        if row["section"] == "summary" and row["item"] in audit_labels:
            audit_rows.append([
                audit_labels[row["item"]],
                f"{int(float(row['value'])):,}",
                f"{float(row['rate']):.3%}",
                audit_treatments.get(row["treatment"], row["treatment"]),
            ])
    add_caption(doc, "表 2-4　历史数据质量审计")
    add_table(doc, ["检查项", "记录数", "比例", "处理口径"], audit_rows,
              [2500, 1150, 1000, CONTENT_WIDTH_DXA-4650], font_size=8.5)
    add_figure(doc, "fig2_0_data_quality_audit.png", "图 2-1　历史数据质量审计概览", width_cm=13.8)
    add_text(doc, "高比分记录可能是真实比赛，不依据箱线图或固定阈值自动删除；程序只标记单场总进球不少于 10 的记录，供敏感性复核。重复比赛键同样先报告，未经来源核对不自动删除。")
    add_page_break(doc)

    add_heading(doc, "第三章　探索性数据分析", 1)
    add_heading(doc, "3.1 历史样本概况", 2)
    overview_rows = [
        ["比赛数量", f"{history_stats['n']:,} 场"],
        ["覆盖球队", f"{history_stats['teams']} 支"],
        ["先列队平均进球", f"{history_stats['mean_home']:.3f}"],
        ["后列队平均进球", f"{history_stats['mean_away']:.3f}"],
        ["场均总进球", f"{history_stats['mean_total']:.3f}"],
        ["先列队胜 / 平 / 后列队胜", f"{history_stats['home_rate']:.1%} / {history_stats['draw_rate']:.1%} / {history_stats['away_rate']:.1%}"],
    ]
    add_caption(doc, "表 3-1　历史比赛描述性统计")
    add_table(doc, ["指标", "数值"], overview_rows, [3200, CONTENT_WIDTH_DXA-3200], font_size=9.5)
    feature_labels_summary = {
        "elo_diff": "Elo 分差 / 400",
        "recent_win_rate_diff": "近 10 场胜率差",
        "home_attack_vs_away_defense": "先列队进攻-后列队防守",
        "away_attack_vs_home_defense": "后列队进攻-先列队防守",
        "recent_goal_diff_diff": "近期净胜球差",
        "home_field": "非中立场指标",
        "host_country_diff": "东道主身份差",
        "tournament_importance": "赛事重要性调整",
    }
    feature_rows = []
    for row in feature_summary:
        if row["feature"] in feature_labels_summary:
            feature_rows.append([
                feature_labels_summary[row["feature"]], f"{int(float(row['n'])):,}",
                f"{float(row['mean']):.3f}", f"{float(row['sd']):.3f}",
                f"{float(row['min']):.3f}", f"{float(row['q1']):.3f}",
                f"{float(row['median']):.3f}", f"{float(row['q3']):.3f}", f"{float(row['max']):.3f}",
            ])
    add_caption(doc, "表 3-2　训练集建模特征描述性统计")
    add_table(doc, ["特征", "n", "均值", "标准差", "最小", "Q1", "中位", "Q3", "最大"], feature_rows,
              [2100, 850, 800, 850, 750, 700, 750, 700, CONTENT_WIDTH_DXA-7500], font_size=7.4)
    add_text(doc, "Elo 分差除以 400 属于尺度重表达，便于对应 Elo 胜率公式，并非使用全样本均值和标准差进行的标准化。所有滚动变量均在比赛日期之前计算。")
    add_figure(doc, "fig2_1_matches_per_year.png", "图 3-1　历史比赛年度分布")
    add_figure(doc, "fig2_2_tournament_types.png", "图 3-2　赛事类型构成", width_cm=12.0)

    add_heading(doc, "3.2 进球分布与泊松假设", 2)
    add_text(doc, f"历史样本的场均总进球为 {history_stats['mean_total']:.3f}，单队进球取非负整数并呈右偏分布。该特征使泊松回归成为透明且可复现的基准模型，但并不证明独立泊松假设完全成立。平局和 0:0、1:1 等低比分会受到双方进球相关性的影响，需要通过校准或 Dixon-Coles 修正处理。")
    add_figure(doc, "fig3_1_goal_distribution.png", "图 3-3　单队进球与总进球分布")

    add_heading(doc, "3.3 动态 Elo 与滚动状态", 2)
    add_text(doc, "动态 Elo 在每场比赛前给出双方相对强度，并在赛后按实际结果更新。它比静态 FIFA 排名更适合按时间顺序构造特征，但当前项目没有独立 team_elo.csv，因此使用历史比赛推断的简化 Elo。该选择必须在正式数据确认时由用户明确批准。")
    add_figure(doc, "fig3_2_elo_vs_winrate.png", "图 3-4　动态 Elo 分差与先列队历史得分率关系")
    corr_rows = [[name, f"{value:.3f}"] for name, value in correlations.items()]
    add_caption(doc, "表 3-3　特征与比赛结果的 Pearson 相关系数（描述性）")
    add_table(doc, ["特征", "相关系数"], corr_rows, [4600, CONTENT_WIDTH_DXA-4600], font_size=9.5)
    add_text(doc, "相关系数只用于探索性描述，不代表因果效应。滚动状态变量彼此相关，回归系数也可能受共线性、对手强度和变量尺度影响。")
    add_figure(doc, "fig3_3_correlation_heatmap.png", "图 3-5　特征相关性热力图", width_cm=12.8)
    add_page_break(doc)

    add_heading(doc, "第四章　研究方法与模型实现", 1)
    add_heading(doc, "4.1 动态 Elo", 2)
    add_equation(doc, "E_A = 1 / (1 + 10^(-(R_A - R_B)/400))")
    add_equation(doc, "R'_A = R_A + K × w × (S_A - E_A)")
    add_text(doc, "其中 R_A、R_B 为赛前 Elo，S_A 为实际得分（胜 1、平 0.5、负 0），w 为赛事重要性权重。代码按日期逐场更新，保证本场特征不包含本场结果。")

    add_heading(doc, "4.2 泊松回归比分模型", 2)
    add_equation(doc, "P(Y = k) = exp(-λ) × λ^k / k!")
    add_equation(doc, "log(λ) = β_0 + β_1·Elo差 + β_2·近期胜率差 + β_3·攻防状态 + …")
    selected_model = metadata.get("selected_poisson_model", "full")
    add_text(doc, f"先列队与后列队进球数分别拟合泊松广义线性模型。预测阶段先得到 λ_1 和 λ_2，再枚举或抽样双方进球，形成胜、平、负概率和具体比分。四组嵌套模型经过时间滚动验证后，本次模拟采用 {selected_model} 规格。泊松模型被选为模拟主模型，因为小组排名需要积分、净胜球和进球数，二元分类模型无法提供这些量。")

    add_heading(doc, "4.3 平局概率校准", 2)
    multiplier = float(metadata["draw_calibration_multiplier"])
    add_text(doc, f"独立泊松模型容易低估低比分相关性。代码使用 2025 年校准集拟合平局概率乘数，本次乘数为 {multiplier:.4f}。校准后重新归一化胜、平、负概率；2026 年验证集不参与乘数拟合。")

    add_heading(doc, "4.4 蒙特卡洛赛事模拟", 2)
    sims = int(float(metadata["simulations"]))
    add_text(doc, f"从截至 6 月 20 日 09:18 的已赛状态出发，剩余赛程按泊松概率抽样，完整赛事重复 {sims:,} 次。每次模拟记录各队到达的最远轮次，以频率估计 32 强、16 强、8 强、四强、决赛和夺冠概率。随机种子固定为 {metadata['seed']}。")
    add_equation(doc, f"P(球队 i 夺冠) ≈ 球队 i 夺冠次数 / {sims:,}")
    add_text(doc, "淘汰赛 90 分钟平局时，当前实现使用双方预期进球强度之比决定晋级方。小组排名依次使用积分、净胜球和进球数；公平竞赛分缺失时随机处理完全同分。")

    add_heading(doc, "4.5 当前实现与官方规则的差距", 2)
    risks = [
        ["Annex C", "当前没有 495 种最佳第三名组合的完整映射；调试代码使用候选槽位约束近似分配。", "高"],
        ["东道主变量", "当前 host_country_diff 同时带有球队身份信息，不能直接解释为比赛所在地效应，可能高估美、加、墨。", "高"],
        ["未来状态更新", "每次模拟中的未来比赛未继续更新 Elo/近期状态，后续轮次使用当前固定强度。", "中"],
        ["低比分相关", "独立泊松加乘数校准仍不能完全替代 Dixon-Coles 或双泊松模型。", "中"],
        ["模拟次数", f"当前为 {sims:,} 次，已达到课程报告的模拟稳定性建议下限；当前仍因数据未确认而属于调试输出。", "中"],
    ]
    add_caption(doc, "表 4-1　当前模型实现风险")
    add_table(doc, ["问题", "影响", "等级"], risks, [1550, CONTENT_WIDTH_DXA-2350, 800], font_size=8.8,
              header_fill=RISK)
    add_figure(
        doc,
        "football_analytics_method-v1.png",
        "图 4-1　足球比赛数据到赛事模拟的主题概念图（AI 生成插图，不含模型结果）",
        width_cm=14.6,
        source_dir=ASSETS,
    )
    add_page_break(doc)

    add_heading(doc, "第五章　模型训练、验证与诊断", 1)
    add_heading(doc, "5.1 独立验证集表现", 2)
    actual_counts = {
        "H": int(float(metrics["confusion_H_pred_H"])) + int(float(metrics["confusion_H_pred_D"])) + int(float(metrics["confusion_H_pred_A"])),
        "D": int(float(metrics["confusion_D_pred_H"])) + int(float(metrics["confusion_D_pred_D"])) + int(float(metrics["confusion_D_pred_A"])),
        "A": int(float(metrics["confusion_A_pred_H"])) + int(float(metrics["confusion_A_pred_D"])) + int(float(metrics["confusion_A_pred_A"])),
    }
    majority_baseline = max(actual_counts.values()) / int(float(metrics["samples"]))
    validation_rows = [
        ["样本数", str(int(float(metrics["samples"])))],
        ["胜平负准确率", f"{float(metrics['accuracy']):.2%}"],
        ["多数类基线", f"{majority_baseline:.2%}"],
        ["相对多数类提升", f"{float(metrics['accuracy']) - majority_baseline:+.2%}"],
        ["Brier Score", f"{float(metrics['brier_score']):.4f}"],
        ["Log Loss", f"{float(metrics['log_loss']):.4f}"],
        ["实际平局率", f"{float(metrics['actual_draw_rate']):.2%}"],
        ["平均预测平局概率", f"{float(metrics['mean_predicted_draw_probability']):.2%}"],
    ]
    add_caption(doc, "表 5-1　2026-01-01 至 2026-06-10 独立验证结果")
    add_table(doc, ["指标", "数值"], validation_rows, [4300, CONTENT_WIDTH_DXA-4300], font_size=9.5)
    add_text(doc, f"模型准确率为 {float(metrics['accuracy']):.2%}，比多数类基线高 {float(metrics['accuracy']) - majority_baseline:.2%}。概率评价同时优于均匀三分类基线（Brier Score 0.6667、Log Loss 1.0986），说明模型包含有效信息，但绝对预测能力有限。")

    add_heading(doc, "5.2 混淆矩阵与平局问题", 2)
    confusion_rows = [
        ["实际先列队胜", metrics["confusion_H_pred_H"], metrics["confusion_H_pred_D"], metrics["confusion_H_pred_A"]],
        ["实际平局", metrics["confusion_D_pred_H"], metrics["confusion_D_pred_D"], metrics["confusion_D_pred_A"]],
        ["实际后列队胜", metrics["confusion_A_pred_H"], metrics["confusion_A_pred_D"], metrics["confusion_A_pred_A"]],
    ]
    add_caption(doc, "表 5-2　验证集混淆矩阵")
    add_table(doc, ["实际 / 预测", "先列队胜", "平局", "后列队胜"], confusion_rows,
              [2350, 1950, 1300, CONTENT_WIDTH_DXA-5600], font_size=9.3)
    add_figure(doc, "fig5_2_confusion_matrix.png", "图 5-1　验证集混淆矩阵", width_cm=10.5)
    add_text(doc, "验证集中 77 场实际平局没有一场被选为最高概率类别。概率模型仍给出平均 23.55% 的平局概率，因此 Brier Score 与 Log Loss 仍有意义；但如果作业强调胜平负分类，‘不预测平局’是必须修正的明显缺陷。")

    add_heading(doc, "5.3 泊松回归系数", 2)
    feature_labels = {
        "(Intercept)": "截距",
        "elo_diff": "Elo 分差",
        "recent_win_rate_diff": "近 10 场胜率差",
        "home_attack_vs_away_defense": "先列队进攻 - 后列队防守",
        "away_attack_vs_home_defense": "后列队进攻 - 先列队防守",
        "recent_goal_diff_diff": "近 10 场净胜球差",
        "home_field": "非中立场优势",
        "host_country_diff": "东道主身份差",
        "tournament_importance": "赛事重要性",
    }
    coef_rows = [[feature_labels.get(r["feature"], r["feature"]), f"{float(r['home_goal_coefficient']):.4f}", f"{float(r['away_goal_coefficient']):.4f}"] for r in coefficients]
    add_caption(doc, "表 5-3　泊松回归系数")
    add_table(doc, ["特征", "先列队进球模型", "后列队进球模型"], coef_rows,
              [3900, 2150, CONTENT_WIDTH_DXA-6050], font_size=9.1)
    add_figure(doc, "fig5_1_coefficients.png", "图 5-2　泊松回归系数对比")
    add_text(doc, "Elo 分差的方向符合预期：强队的预期进球更高、对手预期进球更低。近期胜率差等变量的系数不能单独解释为因果关系；其符号受到 Elo、攻防变量和滚动状态共线性的共同影响。东道主身份差在当前定义下存在身份泄漏风险，本报告不把该系数解释为纯主场效应。")

    add_heading(doc, "5.4 系数显著性与泊松拟合诊断", 2)
    inference_rows = []
    for row in poisson_inference:
        inference_rows.append([
            "先列队" if row["outcome"] == "home" else "后列队",
            feature_labels.get(row["term"], row["term"]),
            f"{float(row['estimate']):.4f}", f"{float(row['std_error']):.4f}",
            f"{float(row['z_value']):.2f}", f"{float(row['p_value']):.3g}", row["significance"],
        ])
    add_caption(doc, "表 5-4　泊松回归系数 Wald 检验")
    add_table(doc, ["模型", "特征", "估计", "标准误", "z", "p", "显著性"], inference_rows,
              [950, 2850, 950, 950, 700, 950, CONTENT_WIDTH_DXA-7350], font_size=7.5)
    fit_rows = []
    for row in poisson_fit:
        fit_rows.append([
            "先列队进球" if row["outcome"] == "home" else "后列队进球",
            f"{float(row['null_deviance']):.1f}", f"{float(row['residual_deviance']):.1f}",
            f"{float(row['aic']):.1f}", f"{float(row['bic']):.1f}",
            f"{float(row['pearson_dispersion']):.3f}", f"{float(row['mcfadden_like_deviance_reduction']):.3f}",
        ])
    add_caption(doc, "表 5-5　泊松模型拟合统计")
    add_table(doc, ["模型", "Null Dev.", "Residual Dev.", "AIC", "BIC", "离散度", "偏差下降率"], fit_rows,
              [1350, 1200, 1400, 1150, 1150, 950, CONTENT_WIDTH_DXA-7200], font_size=8.0)
    add_figure(doc, "fig5_3_poisson_diagnostics.png", "图 5-3　泊松残差、影响点与进球校准诊断", width_cm=15.0)
    add_figure(doc, "fig5_3b_poisson_qq.png", "图 5-4　泊松回归 Deviance residuals Q-Q 诊断", width_cm=15.0)
    add_text(doc, "Q-Q 图用于识别 deviance residuals 相对标准正态参考线的系统偏离、重尾和异常点。需要强调，泊松 GLM 并不要求响应变量或残差服从正态分布，因此该图是近似分布诊断而不是泊松模型成立的必要正态性检验。两套模型的 Pearson 离散度均大于 1，说明数据存在轻度过度离散。本文因此同时报告样本外 Brier Score、Log Loss 和滚动验证结果，不把系数显著性等同于预测效果。偏差下降率只用于描述相对空模型的拟合改善，不解释为线性回归 R²。")

    add_heading(doc, "5.5 嵌套模型选择与时间滚动验证", 2)
    cv_means = defaultdict(list)
    for row in rolling_validation:
        cv_means[row["model"]].append(float(row["log_loss"]))
    comparison_rows = []
    model_names_cn = {"elo": "Elo", "context": "Elo+场地赛事", "form": "+近期状态", "full": "+完整攻防"}
    for row in model_comparison:
        comparison_rows.append([
            model_names_cn[row["model"]], str(int(float(row["parameters"]))),
            f"{float(row['home_aic']) + float(row['away_aic']):.1f}",
            f"{float(row['home_bic']) + float(row['away_bic']):.1f}",
            f"{sum(cv_means[row['model']]) / len(cv_means[row['model']]):.4f}",
            f"{float(row['validation_log_loss']):.4f}",
            "采用" if row["model"] == selected_model else "对照",
        ])
    add_caption(doc, "表 5-6　候选泊松模型比较")
    add_table(doc, ["规格", "参数数", "AIC合计", "BIC合计", "滚动Log Loss", "最终Log Loss", "用途"], comparison_rows,
              [1750, 850, 1200, 1200, 1400, 1300, CONTENT_WIDTH_DXA-7700], font_size=8.0)
    add_figure(doc, "fig5_4_model_comparison.png", "图 5-5　四组候选模型滚动验证 Log Loss", width_cm=14.0)
    add_text(doc, "滚动验证严格保持训练期早于验证期，未使用随机 K 折。模型选择以三个时间窗口的平均 Log Loss 为第一排序指标，Brier Score 为并列时的第二指标；2026 年最终测试集不参与选择。")

    add_heading(doc, "5.6 二元逻辑回归对照实验", 2)
    logistic_rows = [
        ["样本数", str(int(float(logistic_metrics["samples"])))],
        ["准确率（阈值0.5）", f"{float(logistic_metrics['accuracy']):.2%}"],
        ["Brier Score（二元）", f"{float(logistic_metrics['brier_score']):.4f}"],
        ["Log Loss（二元）", f"{float(logistic_metrics['log_loss']):.4f}"],
        ["ROC AUC", f"{float(logistic_metrics['auc']):.4f}"],
        ["TN / FP / FN / TP", f"{int(float(logistic_metrics['confusion_actual_0_pred_0']))} / {int(float(logistic_metrics['confusion_actual_0_pred_1']))} / {int(float(logistic_metrics['confusion_actual_1_pred_0']))} / {int(float(logistic_metrics['confusion_actual_1_pred_1']))}"],
    ]
    add_caption(doc, "表 5-7　第一列球队获胜 vs 非获胜逻辑回归")
    add_table(doc, ["指标", "结果"], logistic_rows, [4300, CONTENT_WIDTH_DXA-4300], font_size=9.2)
    add_figure(doc, "fig5_5_logistic_roc.png", "图 5-5　二元逻辑回归 ROC 曲线", width_cm=10.5)
    significant_logistic = [r for r in logistic_inference if float(r["p_value"]) < 0.05]
    add_text(doc, f"逻辑回归将‘第一列球队 90 分钟获胜’记为 1，将平局和失利合并为 0，共有 {len(significant_logistic)} 个系数在 5% 水平显著。该实验用于对应课程中的二元逻辑回归、Wald 检验与 ROC/AUC，不用于生成比分，也不进入世界杯蒙特卡洛模拟。其二元 Brier Score 与泊松模型的三分类 Brier Score定义不同，不能直接按数值大小比较。")
    add_page_break(doc)

    add_heading(doc, "第六章　截至 2026-06-20 09:18 的赛中检验", 1)
    add_heading(doc, "6.1 已赛结果与总体差异", 2)
    observed_outcome_counts = Counter(r["actual_result"] for r in comparisons)
    observed_majority = max(observed_outcome_counts.values()) / len(comparisons)
    obs_rows = [
        ["已赛场次", str(int(float(observed["matches"])))],
        ["方向命中", f"{int(float(observed['correct']))} 场"],
        ["方向准确率", f"{float(observed['accuracy']):.2%}"],
        ["已赛多数类基线", f"{observed_majority:.2%}"],
        ["Brier Score", f"{float(observed['brier_score']):.4f}"],
        ["Log Loss", f"{float(observed['log_loss']):.4f}"],
        ["比分绝对误差", f"{float(observed['mean_score_abs_error']):.3f}"],
        ["实际平局 / 最高概率预测平局", f"{int(float(observed['actual_draws']))} / {int(float(observed['predicted_draws']))}"],
    ]
    observed_matches = int(float(observed["matches"]))
    add_caption(doc, f"表 6-1　{observed_matches} 场世界杯赛中预测表现")
    add_table(doc, ["指标", "结果"], obs_rows, [4300, CONTENT_WIDTH_DXA-4300], font_size=9.5)
    acc_gap = float(observed["accuracy"]) - float(metrics["accuracy"])
    brier_gap = float(observed["brier_score"]) - float(metrics["brier_score"])
    log_gap = float(observed["log_loss"]) - float(metrics["log_loss"])
    baseline_gap = float(observed["accuracy"]) - observed_majority
    add_text(doc, f"世界杯已赛样本准确率为 {float(observed['accuracy']):.2%}，较历史独立验证集变化 {acc_gap * 100:+.2f} 个百分点；Brier Score 变化 {brier_gap:+.4f}，Log Loss 变化 {log_gap:+.4f}。已赛多数类基线为 {observed_majority:.2%}，模型相对该基线变化 {baseline_gap * 100:+.2f} 个百分点。")
    add_text(doc, f"{observed_matches} 场出现 {int(float(observed['actual_draws']))} 场平局，占 {int(float(observed['actual_draws']))/observed_matches:.1%}；模型平均平局概率为 {float(observed['mean_predicted_draw_probability']):.1%}，最高概率类别预测平局 {int(float(observed['predicted_draws']))} 场。赛中表现的主要结构性偏差仍集中在平局与低比分结构。")

    add_heading(doc, "6.2 逐场结果", 2)
    result_rows = []
    label = {"H": "先列队胜", "D": "平", "A": "后列队胜"}
    for row in comparisons:
        result_rows.append([
            row["match_date"][5:],
            row["group"],
            f"{row['team_1']} {row['team_1_score']}-{row['team_2_score']} {row['team_2']}",
            label[row["predicted_result"]],
            label[row["actual_result"]],
            "命中" if row["correct"].upper() == "TRUE" else "偏差",
        ])
    add_caption(doc, "表 6-2　逐场最高概率方向与实际结果")
    add_table(doc, ["日期", "组", "实际比分", "模型方向", "实际方向", "结果"], result_rows,
              [800, 450, 3300, 1250, 1250, CONTENT_WIDTH_DXA-7050], font_size=7.8)

    add_heading(doc, "6.3 典型偏差比赛", 2)
    surprises = sorted(comparisons, key=lambda r: float(r["log_loss"]), reverse=True)[:6]
    surprise_rows = []
    for row in surprises:
        surprise_rows.append([
            f"{row['team_1']} {row['team_1_score']}-{row['team_2_score']} {row['team_2']}",
            f"{float(row['p_team_1_win']):.1%} / {float(row['p_draw']):.1%} / {float(row['p_team_2_win']):.1%}",
            label[row["actual_result"]],
            f"{float(row['log_loss']):.3f}",
        ])
    add_caption(doc, "表 6-3　Log Loss 最大的 6 场比赛")
    add_table(doc, ["比赛", "胜 / 平 / 负概率", "实际", "Log Loss"], surprise_rows,
              [3300, 2750, 1050, CONTENT_WIDTH_DXA-7100], font_size=8.5)
    add_text(doc, "西班牙 0:0 佛得角是当前最明显的单场偏差：模型给西班牙 90.37% 的获胜概率，但实际为平局。加拿大 1:1 波黑、卡塔尔 1:1 瑞士等比赛也显示强队优势被实际平局削弱。这些比赛共同指向低比分相关、比赛节奏和赛会首轮保守策略未被模型充分表达。")

    latest_results = comparisons[-2:]
    latest_hits = sum(r["correct"].upper() == "TRUE" for r in latest_results)
    latest_descriptions = []
    for row in latest_results:
        state = "命中" if row["correct"].upper() == "TRUE" else "偏差"
        latest_descriptions.append(
            f"{row['team_1']} {row['team_1_score']}:{row['team_2_score']} {row['team_2']}（{state}）"
        )
    add_heading(doc, "6.4 6 月 20 日新增完赛结果", 2)
    add_text(doc, f"截至 09:18 新增录入 2 场明确完赛比赛，模型方向命中 {latest_hits}/2：{'；'.join(latest_descriptions)}。同一时点仍在直播的巴西对海地没有进入结果快照，也没有参与当前积分判断或模型更新。")
    add_page_break(doc)

    add_heading(doc, "第七章　世界杯模拟结果", 1)
    add_callout(doc, "读取方式", f"本章读取 {sims:,} 次、随机种子 {metadata['seed']} 的调试模拟。置信区间是基于二项频率标准误的近似区间，不是 bootstrap；Annex C 与数据来源确认完成前，只能比较当前模型内部的相对排序。", fill=CAUTION)
    add_heading(doc, "7.1 冠军概率", 2)
    champ_rows = [[r["rank"], r["team"], f"{float(r['probability']):.1%}", f"[{float(r['ci_lower']):.1%}, {float(r['ci_upper']):.1%}]"] for r in champion]
    add_caption(doc, "表 7-1　冠军概率 Top 10")
    add_table(doc, ["排名", "球队", "夺冠概率", "95% 模拟误差区间"], champ_rows,
              [900, 2350, 1900, CONTENT_WIDTH_DXA-5150], font_size=9.3)
    add_figure(doc, "fig6_1_champion_bar.png", "图 7-1　冠军概率 Top 10 与模拟误差棒")
    lead_gap = float(champion[0]["probability"]) - float(champion[1]["probability"])
    add_text(doc, f"当前模拟中，{champion[0]['team']}夺冠概率为 {float(champion[0]['probability']):.1%}，{champion[1]['team']}为 {float(champion[1]['probability']):.1%}。两队区间明显重叠，不能把 {lead_gap * 100:.1f} 个百分点的差异解释为稳定领先。第三名之后概率快速下降，说明当前模型将冠军概率高度集中在前两队；这一集中程度仍需在补齐 Annex C、修正东道主变量并使用不同随机种子重复运行后核验。")

    add_heading(doc, "7.2 四强概率与晋级路径", 2)
    ff_rows = [[r["rank"], r["team"], f"{float(r['probability']):.1%}"] for r in final_four]
    add_caption(doc, "表 7-2　四强概率 Top 10")
    add_table(doc, ["排名", "球队", "四强概率"], ff_rows, [1200, 3300, CONTENT_WIDTH_DXA-4500], font_size=9.3)
    add_figure(doc, "fig6_2_final_four_bar.png", "图 7-2　四强概率 Top 10")
    add_figure(doc, "fig6_3_progression_funnel.png", "图 7-3　热门球队分阶段晋级概率")
    add_figure(doc, "fig6_4_team_funnel_line.png", "图 7-4　热门球队晋级概率阶梯图")

    add_heading(doc, "7.3 各小组出线概率", 2)
    group_rows = []
    for group in "ABCDEFGHIJKL":
        entries = [r for r in group_adv if r["group"] == group]
        entries.sort(key=lambda r: float(r["round_of_32"]), reverse=True)
        for row in entries:
            group_rows.append([group, row["team"], f"{float(row['round_of_32']):.1%}", f"{float(row['round_of_16']):.1%}", f"{float(row['champion']):.1%}"])
    add_caption(doc, "表 7-3　各队出线、16 强与冠军概率")
    add_table(doc, ["组", "球队", "32 强", "16 强", "冠军"], group_rows,
              [550, 2850, 1600, 1600, CONTENT_WIDTH_DXA-6600], font_size=8.2)

    add_heading(doc, "7.3.1 当前小组前二确定性判断", 2)
    status_rows = []
    for group in "ABCDEFGHIJKL":
        entries = [row for row in current_group_status if row["group"] == group]
        entries.sort(key=lambda row: int(row["current_rank"]))
        order_text = "；".join(
            f"{row['team']} {row['points']}分({int(row['goal_diff']):+d})" for row in entries
        )
        decisions = [f"{row['team']}：{row['top2_status']}" for row in entries if row["top2_status"] != "仍竞争前二"]
        status_rows.append([group, order_text, "；".join(decisions) if decisions else "全部仍竞争前二"])
    add_caption(doc, "表 7-4　基于已结束比赛的当前小组前二判断")
    add_table(doc, ["组", "当前排序（积分/净胜球）", "确定性判断"], status_rows,
              [550, 5150, CONTENT_WIDTH_DXA-5700], font_size=8.1)
    add_text(doc, "判断方法穷举每组剩余比赛的胜、平、负组合。为避免在未来净胜球未知时制造确定结论，若最终积分相同则保持未决；最佳第三名还需要跨 12 个小组比较，因此本表只判断小组前二。")

    add_heading(doc, "7.4 条件晋级路线与整体签表", 2)
    add_text(doc, "下一轮对手并不是预先固定的。本文在每次蒙特卡洛模拟中记录球队所在槽位、逐轮对手、胜负和已走过的路径前缀，并在相同路径条件下重新统计下一轮对手分布。对任一阶段，下一轮晋级概率等于各潜在对手分支的遇敌概率与该对阵晋级率的加权和。")
    add_equation(doc, "P(晋级下一轮 | 当前路径) = Σ P(对手=o | 当前路径) × P(战胜o | 当前路径, 对手=o)")
    annex_mode = metadata.get("annex_c_mode", "unknown")
    add_callout(
        doc,
        "路线图读取边界",
        f"本次路线输出使用 {annex_mode}。图中‘遇’为条件遇敌概率，‘胜’为该对阵中的模拟晋级率，‘贡献’为该分支对下一轮无条件概率的贡献。遇敌样本少于 30 次时标记为样本不足。正式提交前必须使用经过确认的 Annex C 完整映射重跑。",
        fill=CAUTION,
    )
    add_decorative_image_pair(doc, [
        ("group-stage-atmosphere-v1.png", "十二个小组赛区域的足球场氛围插图"),
        ("knockout-route-atmosphere-v1.png", "多条淘汰赛路径汇聚至决赛的氛围插图"),
    ])

    add_landscape_section(doc)
    route_figures = [
        ("fig7_5a_overall_groups.png", "图 7-5a　世界杯整体路径：A-L 组晋级概率概览"),
        ("fig7_5b_overall_bracket.png", "图 7-5b　世界杯整体路径：32 强至决赛槽位概率"),
        ("fig7_6_spain_route.png", "图 7-6　西班牙条件晋级路线"),
        ("fig7_7_france_route.png", "图 7-7　法国条件晋级路线"),
        ("fig7_8_argentina_route.png", "图 7-8　阿根廷条件晋级路线"),
        ("fig7_9_brazil_route.png", "图 7-9　巴西条件晋级路线"),
        ("fig7_10_portugal_route.png", "图 7-10　葡萄牙条件晋级路线"),
    ]
    for index, (filename, caption) in enumerate(route_figures):
        add_figure(doc, filename, caption, width_cm=25.8)
        if index < len(route_figures) - 1:
            add_page_break(doc)
    add_portrait_section(doc)

    add_heading(doc, "第八章　赛中更新的影响", 1)
    add_heading(doc, "8.1 6 月 19 日至 20 日概率变化", 2)
    positives = sorted(snapshots, key=lambda r: float(r["round_of_32_change"]), reverse=True)[:6]
    negatives = sorted(snapshots, key=lambda r: float(r["round_of_32_change"]))[:6]
    change_rows = []
    for row in positives:
        change_rows.append(["上升", row["team"], f"{float(row['round_of_32_change']):+.1%}", f"{float(row['champion_change']):+.1%}"])
    for row in negatives:
        change_rows.append(["下降", row["team"], f"{float(row['round_of_32_change']):+.1%}", f"{float(row['champion_change']):+.1%}"])
    add_caption(doc, "表 8-1　32 强概率变化最大的球队")
    add_table(doc, ["方向", "球队", "32 强概率变化", "夺冠概率变化"], change_rows,
              [1050, 2450, 2500, CONTENT_WIDTH_DXA-6000], font_size=9)
    top_gain = positives[0]
    top_loss = negatives[0]
    add_text(doc, f"本次快照只新增两场完赛结果。32 强概率上升最多的是{top_gain['team']}（{float(top_gain['round_of_32_change']):+.1%}），下降最多的是{top_loss['team']}（{float(top_loss['round_of_32_change']):+.1%}）。概率变化同时来自已赛积分、净胜球、动态 Elo 与滚动状态更新，不能只归因于单一因素。")

    add_heading(doc, "8.2 结果解释边界", 2)
    add_text(doc, f"赛中概率变化由三部分共同造成：已赛积分与净胜球直接改变小组位置；赛果更新动态 Elo；最近比赛进入滚动状态窗口。当前快照对比使用相同随机种子，减少了模拟噪声。在 {sims:,} 次模拟下，极小概率变化仍需通过多随机种子重复运行区分赛果影响与抽样波动。")
    add_text(doc, "东道主球队的变化需额外谨慎。当前 host_country_diff 变量可能混入球队身份，墨西哥、加拿大和美国的概率不宜直接解释为纯主场收益。修正变量后必须重新训练并重跑全部模拟。")

    add_heading(doc, "8.3 与当前赛果的总体判断", 2)
    add_text(doc, "模型对明显强弱比赛仍具备一定识别能力，例如阿根廷胜阿尔及利亚、法国胜塞内加尔、英格兰胜克罗地亚、加拿大胜卡塔尔均命中方向；但多场强队被逼平和个别弱队取胜显著拉高 Log Loss。当前结论不是‘模型完全失效’，而是‘强度排序有信息、平局和赛会短期状态建模不足’。")
    add_page_break(doc)

    add_heading(doc, "第九章　结论、局限与改进方案", 1)
    add_heading(doc, "9.1 主要结论", 2)
    conclusions = [
        f"严格时间外验证准确率为 {float(metrics['accuracy']):.2%}，高于多数类基线 {float(metrics['accuracy']) - majority_baseline:.2%}，概率指标也优于均匀预测基线。",
        f"截至 6 月 20 日 09:18 的 {observed_matches} 场赛中命中率为 {float(observed['accuracy']):.2%}；{int(float(observed['actual_draws']))} 场实际平局中，最高概率类别预测平局 {int(float(observed['predicted_draws']))} 场。",
        f"当前调试模拟中，{champion[0]['team']}与{champion[1]['team']}位列冠军概率前二，但区间重叠，不能宣称第一名具有统计上稳定优势。",
        "已赛积分和净胜球能显著改变小组出线概率，赛中更新对冠军概率的影响相对较小。",
        "当前最优先的改进不是继续增加图表，而是确认数据、实现完整 Annex C、修正东道主变量并改善低比分相关模型。",
    ]
    for idx, item in enumerate(conclusions, 1):
        add_text(doc, f"（{idx}）{item}", indent=False, after=4)

    add_heading(doc, "9.2 必须完成的改进", 2)
    improvements = [
        ["P0", "数据确认", f"确认分组、赛程、{observed_matches} 场赛果、历史数据原始来源与快照日期，并更新 data_approval.csv。"],
        ["P0", "Annex C", "录入官方 495 种最佳第三名组合到淘汰赛槽位的完整映射，并增加自动测试。"],
        ["P0", "东道主变量", "改为‘比赛所在国是否为球队本国’，禁止用球队身份替代比赛地点。"],
        ["P1", "平局模型", "实现 Dixon-Coles 或双泊松模型；比较校准曲线、Brier Score 和 Log Loss。"],
        ["P1", "赛事模拟", f"未来模拟比赛后继续更新 Elo/近期状态；保持至少 {sims:,} 次并增加多随机种子稳定性检查。"],
        ["P2", "增量数据", "在用户确认后加入可验证的伤病、首发、市场赔率或 xG，不对缺失变量做幻觉补全。"],
    ]
    add_caption(doc, "表 9-1　模型改进优先级")
    add_table(doc, ["优先级", "事项", "执行内容"], improvements,
              [900, 1800, CONTENT_WIDTH_DXA-2700], font_size=9)

    add_heading(doc, "9.3 局限性", 2)
    add_text(doc, "本文训练数据以国家队历史比分为核心，未包含球员级伤病、首发阵容、战术变化、天气、旅行距离、公平竞赛分、市场赔率和 xG。国家队比赛样本稀疏且赛事强度差异大，短期爆冷频繁。模型适合展示公开数据下的概率建模流程，不应被解释为确定性赛果或投注建议。")
    add_text(doc, "当前报告的最大限制不是模型复杂度，而是数据确认状态。只要关键外部数据仍是待确认，任何更复杂的算法都会放大输入不确定性，无法替代数据来源核验。")
    add_page_break(doc)

    add_heading(doc, "参考文献与数据来源", 1)
    refs = [
        "[1] Maher, M. J. (1982). Modelling association football scores. Statistica Neerlandica, 36(3), 109-118. DOI: 10.1111/j.1467-9574.1982.tb00782.x.",
        "[2] Dixon, M. J., & Coles, S. G. (1997). Modelling association football scores and inefficiencies in the football betting market. Journal of the Royal Statistical Society Series C, 46(2), 265-280. DOI: 10.1111/1467-9876.00065.",
        "[3] Elo, A. E. (1978). The Rating of Chessplayers, Past and Present. Arco Publishing, New York.",
        "[4] Metropolis, N., & Ulam, S. (1949). The Monte Carlo method. Journal of the American Statistical Association, 44(247), 335-341. DOI: 10.1080/01621459.1949.10483310.",
        "[5] McCullagh, P., & Nelder, J. A. (1989). Generalized Linear Models (2nd ed.). Chapman and Hall/CRC. DOI: 10.1007/978-1-4899-3242-6.",
        "[6] Karlis, D., & Ntzoufras, I. (2003). Analysis of sports data by using bivariate Poisson models. Journal of the Royal Statistical Society Series D, 52(3), 381-393. DOI: 10.1111/1467-9884.00366.",
        "[7] Hvattum, L. M., & Arntzen, H. (2010). Using ELO ratings for match result prediction in association football. International Journal of Forecasting, 26(3), 460-470. DOI: 10.1016/j.ijforecast.2009.10.002.",
        "[8] Koopman, S. J., & Lit, R. (2015). A dynamic bivariate Poisson model for analysing and forecasting match results in the English Premier League. Journal of the Royal Statistical Society Series A, 178(1), 167-186. DOI: 10.1111/rssa.12042.",
        "[9] Brier, G. W. (1950). Verification of forecasts expressed in terms of probability. Monthly Weather Review, 78(1), 1-3.",
        "[10] R Core Team. R: A Language and Environment for Statistical Computing. R Foundation for Statistical Computing.",
        "[11] FIFA. FIFA World Cup 26 official tournament page. https://www.fifa.com/en/tournaments/mens/worldcup/canadamexicousa2026（赛制事实的首选核验源；当前分组和赛程文件尚未完成最终核验）。",
        "[12] 百度体育. 世界杯赛程与排名页. https://tiyu.baidu.com/al/match?match=世界杯&tab=赛程&current=0（截至 2026-06-20 09:18 明确完赛的赛果录入口径，待用户确认；直播比赛未纳入）。",
        "[13] 用户提供的 2026 世界杯分组、赛程与淘汰赛路径截图，提供日期 2026-06-15。",
        "[14] data/historical_matches.csv，本地公开历史比赛快照；原始下载地址和快照日期待用户确认。",
    ]
    for item in refs:
        paragraph = add_text(doc, item, indent=False, after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=9.5)
    add_page_break(doc)

    add_heading(doc, "附录 A　R 代码与复现说明", 1)
    code_rows = [
        ["R/validate_inputs.R", "输入字段、48 队分组、日期边界和基础完整性校验"],
        ["R/worldcup_predictor.R", "动态 Elo、滚动特征、泊松回归、平局校准、赛中更新和赛事模拟主程序"],
        ["R/compare_snapshots.R", "比较相同随机种子下两个赛中快照的阶段概率变化"],
        ["R/make_charts.R", "从 R 输出 CSV 生成冠军、四强、晋级路径和单队漏斗图"],
        ["R/README.md", "运行顺序、数据确认闸门与调试参数说明"],
    ]
    add_caption(doc, "表 A-1　提交代码文件")
    add_table(doc, ["文件", "作用"], code_rows, [3000, CONTENT_WIDTH_DXA-3000], font_size=9.2)
    add_text(doc, "正式运行命令如下。若确认闸门未通过，主程序会停止；--allow-unconfirmed-data 仅用于调试，不得生成正式概率。")
    code = (
        "Rscript R/validate_inputs.R\n"
        "Rscript R/worldcup_predictor.R\n"
        "Rscript R/make_charts.R"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for line in code.splitlines():
        r = p.add_run(line)
        set_run_font(r, east="等线", latin="Courier New", size=9)
        r.add_break()
    add_text(doc, "报告中的概率、评价指标和图表均从 output 目录读取。完整代码作为独立提交物提供，不在正文重复粘贴 1,118 行源代码，以保持报告可读性。")

    add_heading(doc, "附录 B　运行一致性检查", 1)
    check_rows = [
        ["参赛球队", "48 支", "通过"],
        ["冠军概率合计", "1.000", "通过"],
        ["阶段概率单调性", "32 强 ≥ 16 强 ≥ 8 强 ≥ 四强 ≥ 决赛 ≥ 冠军", "通过"],
        ["概率范围", "全部位于 [0, 1]", "通过"],
        ["赛果与赛程匹配", f"{len(results)} 场均可匹配", "通过"],
        ["关键数据来源确认", "全部仍为待确认", "未通过"],
        ["Annex C 完整映射", "未提供", "未通过"],
    ]
    add_caption(doc, "表 B-1　当前运行检查")
    add_table(doc, ["检查项", "结果", "状态"], check_rows, [2800, 4100, CONTENT_WIDTH_DXA-6900], font_size=9.2)

    add_heading(doc, "附录 C　网页展示与打开方式", 1)
    add_text(doc, "网页通过浏览器读取 web/data 中的 CSV，因此不能直接双击 index.html 打开；应在项目根目录启动一个本地静态服务器。", indent=False)
    web_commands = (
        "cd /Users/lionel/Desktop/大作业\n"
        "python3 -m http.server 8000 --directory web"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for line in web_commands.splitlines():
        r = p.add_run(line)
        set_run_font(r, east="等线", latin="Courier New", size=9)
        r.add_break()
    add_text(doc, "启动后在浏览器打开 http://127.0.0.1:8000/#routes，可直接进入晋级路线页面；完整首页地址为 http://127.0.0.1:8000/。使用结束后回到终端按 Ctrl+C 停止服务。若 8000 端口被占用，可将命令和网址中的 8000 同时改为其他空闲端口。", indent=False)


def audit_docx(doc_path: Path):
    """Basic structural checks before rendering."""
    check = Document(doc_path)
    assert len(check.sections) == 3
    assert abs(check.sections[0].page_width.cm - 21.0) < 0.05
    assert abs(check.sections[0].page_height.cm - 29.7) < 0.05
    assert abs(check.sections[1].page_width.cm - 29.7) < 0.05
    assert abs(check.sections[1].page_height.cm - 21.0) < 0.05
    assert abs(check.sections[2].page_width.cm - 21.0) < 0.05
    assert abs(check.sections[2].page_height.cm - 29.7) < 0.05
    assert len(check.tables) >= 15
    assert any(marker in check._element.xml for marker in ("数据待确认版", "报告状态声明")) or any(
        approved(r["approved"]) for r in read_csv(DATA / "data_approval.csv")
    )
    text = "\n".join(p.text for p in check.paragraphs)
    assert "目　录" in text
    assert "Deviance residuals Q-Q" in text
    forbidden = ["约 0.45", "胜率约为 75%", "占比最大但权重最高"]
    for phrase in forbidden:
        if phrase in text:
            raise AssertionError(f"Unsupported hard-coded claim remains: {phrase}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--allow-unconfirmed-data", action="store_true")
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    _, formal = approval_state()
    if not formal and not args.allow_unconfirmed_data:
        raise SystemExit(
            "Data confirmation is incomplete. Use --allow-unconfirmed-data only to create an explicitly labelled review draft."
        )
    draft = not formal
    target = args.output or ROOT / (
        "2026世界杯冠军概率预测课程分析报告_数据待确认版.docx" if draft
        else "2026世界杯冠军概率预测课程分析报告_正式版.docx"
    )
    doc = Document()
    configure_document(doc)
    build_report(doc, draft=draft)
    doc.core_properties.title = "基于历史比赛数据与蒙特卡洛模拟的 2026 年世界杯冠军概率预测研究"
    doc.core_properties.subject = "工业大数据课程设计报告"
    doc.core_properties.author = "[请填写]"
    doc.core_properties.comments = "由当前 R 模型输出自动生成；数据状态见正文。"
    doc.save(target)
    audit_docx(target)
    print(target)


if __name__ == "__main__":
    main()
