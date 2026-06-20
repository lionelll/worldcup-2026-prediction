#!/usr/bin/env python3
"""Generate the Word report for the 2026 World Cup prediction project."""

import csv
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "output")
FIG_DIR = os.path.join(OUTPUT_DIR, "figures")
APPROVAL_PATH = os.path.join(PROJECT_ROOT, "data", "data_approval.csv")

def approved_value(value):
    return str(value).strip().lower() in {"true", "1", "yes", "y", "approved"}

def enforce_data_approval():
    if "--allow-unconfirmed-data" in sys.argv:
        print("WARNING: --allow-unconfirmed-data enabled. Generated report is for debugging only.")
        return
    required = {
        "worldcup_2026_groups.csv",
        "worldcup_2026_schedule.csv",
        "worldcup_2026_results_asof_2026-06-20.csv",
        "historical_matches.csv",
        "annex_c_full_mapping",
    }
    if not os.path.exists(APPROVAL_PATH):
        raise SystemExit(f"Data approval file not found: {APPROVAL_PATH}")
    with open(APPROVAL_PATH, encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    approval = {row.get("file", ""): approved_value(row.get("approved", "")) for row in rows}
    missing = sorted(required - set(approval))
    pending = sorted(name for name in required if not approval.get(name, False))
    if missing or pending:
        blocked = missing + pending
        raise SystemExit(
            "Data source confirmation required before formal report generation: "
            + ", ".join(blocked)
            + ". Use --allow-unconfirmed-data only for debugging."
        )

def read_csv_data(filename):
    path = os.path.join(OUTPUT_DIR, filename)
    with open(path, encoding="utf-8") as f:
        return list(csv.DictReader(f))

def set_cell_font(cell, text, font_name="宋体", size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.bold = bold

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_font(table.rows[r_idx + 1].cells[c_idx], val)
    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_figure(doc, filename, caption_text, width=Inches(5.5)):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        add_body(doc, f"[图片缺失: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    add_caption(doc, caption_text)

def add_code_block(doc, code_text, title=None):
    if title:
        add_heading(doc, title, level=3)
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.size = Pt(8)
        run.font.name = "Courier New"

def main():
    enforce_data_approval()
    run_metadata = {r["key"]: r["value"] for r in read_csv_data("run_metadata.csv")}
    simulation_count = int(float(run_metadata["simulations"]))
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ===================== COVER PAGE =====================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于历史比赛数据与蒙特卡洛模拟的\n2026 年世界杯冠军概率预测研究")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("工业大数据课程设计报告")
    run.font.size = Pt(16)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for _ in range(3):
        doc.add_paragraph()

    info_lines = [
        "院　系：工学院",
        "专　业：工程管理",
        "小　组：第 X 小组",
        "组　员：姓名 学号",
        "　　　　姓名 学号",
        "　　　　姓名 学号",
        "指导教师：吴建国",
        "日　期：2026 年 6 月",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    # ===================== TABLE OF CONTENTS (placeholder) =====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目　录")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    toc_items = [
        ("第一章　绪论", 1),
        ("　1.1 研究背景", 2),
        ("　1.2 研究目标", 2),
        ("　1.3 数据口径说明", 2),
        ("第二章　数据来源与预处理", 1),
        ("　2.1 数据来源", 2),
        ("　2.2 数据预处理", 2),
        ("　2.3 数据基本情况", 2),
        ("第三章　探索性分析", 1),
        ("　3.1 进球数分布", 2),
        ("　3.2 Elo 分差与胜率", 2),
        ("　3.3 变量相关性", 2),
        ("第四章　研究方法", 1),
        ("　4.1 泊松回归模型", 2),
        ("　4.2 二元逻辑回归对照模型", 2),
        ("　4.3 蒙特卡洛模拟", 2),
        ("第五章　模型训练与评价", 1),
        ("　5.1 模型训练", 2),
        ("　5.2 模型评价指标", 2),
        ("　5.3 模型系数分析", 2),
        ("第六章　预测结果分析", 1),
        ("　6.1 冠军概率预测", 2),
        ("　6.2 四强概率预测", 2),
        ("　6.3 各小组出线概率", 2),
        ("　6.4 赛中结果对照", 2),
        ("第七章　结论与不足", 1),
        ("参考文献", 1),
        ("附录：核心 R 代码", 1),
    ]
    for item, level in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12 if level == 1 else 11)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if level == 1:
            run.font.bold = True

    doc.add_page_break()

    # ===================== CHAPTER 1 =====================
    add_heading(doc, "第一章　绪论", level=1)

    add_heading(doc, "1.1 研究背景", level=2)
    add_body(doc, "世界杯是世界范围内影响力最高的足球赛事之一，比赛结果不仅受到球队整体实力影响，也受到赛制、赛程、对阵路径、临场状态和随机因素影响。传统的赛前预测常以专家判断、赔率或简单排名为依据，但这些方法难以清晰说明不同因素如何共同作用。随着公开体育数据和机器学习方法的发展，利用历史比赛数据建立概率预测模型，已经成为体育数据分析中的常见方法。")
    add_body(doc, "2026 年世界杯由美国、加拿大、墨西哥共同举办，并采用 48 支球队参赛的新赛制。扩军后的赛制增加了小组第三名晋级和 32 强淘汰赛环节，使赛事路径更复杂，也使单纯依赖球队排名进行预测的解释力不足。因此，本文希望通过\u201c单场进球预测模型 + 真实赛制模拟\u201d的方式，对 2026 年世界杯的晋级概率和冠军概率进行量化研究。")

    add_heading(doc, "1.2 研究目标", level=2)
    add_body(doc, "本文的研究目标包括：（1）基于历史国家队比赛结果，构建单场比赛预测模型；（2）使用真实分组、真实赛程和真实淘汰赛路径，模拟 2026 年世界杯完整赛事过程；（3）统计各队小组出线、16 强、8 强、四强、决赛和冠军概率；（4）分析影响球队晋级概率的主要因素，并说明模型的局限性。")

    add_heading(doc, "1.3 数据口径说明", level=2)
    add_body(doc, "本文采用经确认的分组与赛程作为赛事模拟框架。历史模型仅使用 2026 年 6 月 10 日前的比赛信息；赛中更新版本在此基础上纳入截至 2026 年 6 月 20 日 09:18 明确完赛的比分。已赛结果用于当前积分并滚动更新 Elo 和近 10 场状态，不回填训练集，避免把已赛信息误写为赛前预测能力。")

    doc.add_page_break()

    # ===================== CHAPTER 2 =====================
    add_heading(doc, "第二章　数据来源与预处理", level=1)

    add_heading(doc, "2.1 数据来源", level=2)
    add_body(doc, "本文使用的数据包括以下几部分：")
    add_body(doc, "（1）历史比赛数据：来源于公开数据集 International Football Results from 1872 to Present，包含国际足球比赛的日期、双方球队、比分、赛事类型和是否中立场等信息。本文选取 2010 年 1 月 1 日至 2026 年 6 月 10 日的比赛记录，共计 15,817 场。")
    add_body(doc, "（2）2026 年世界杯分组与赛程数据：来源于 FIFA 官方赛事页面和百度体育赛程页，包含 48 支参赛队伍的分组信息（12 个小组、每组 4 队）和完整赛程（72 场小组赛 + 32 场淘汰赛，共 104 场）。")
    add_body(doc, "（3）赛中已赛结果数据：截至 2026 年 6 月 20 日 09:18 明确完成的 30 场小组赛比分；正式报告仅在数据确认闸门通过后使用。")
    add_body(doc, "（4）球队 Elo 评分：由历史比赛结果自动推算。")

    add_heading(doc, "2.2 数据预处理", level=2)
    add_body(doc, "数据预处理包括以下步骤：")
    add_body(doc, "（1）队名统一：不同数据源存在队名差异（如 Korea Republic / South Korea、Côte d'Ivoire / Ivory Coast 等），建立映射表将所有球队名称统一为标准中文名。")
    add_body(doc, "（2）日期截断：赛前训练数据严格截止到 2026 年 6 月 10 日，2026 年 6 月 11 日及之后的比赛结果全部删除，不进入训练集。")
    add_body(doc, "（3）赛事类型权重：不同赛事类型赋予不同权重——世界杯正赛 1.45、洲际杯赛 1.25、预选赛/国家联赛 1.10、友谊赛 0.65、其他 1.0。")
    add_body(doc, "（4）缺失值处理：公开数据中没有的变量（伤病、首发、天气、赔率、xG 等）不强行补造，在局限性部分说明。")

    add_heading(doc, "2.3 数据基本情况", level=2)
    add_body(doc, "历史比赛数据集的基本情况如表 2-1 所示。")

    add_caption(doc, "表 2-1：历史比赛数据集概况")
    add_table(doc,
        ["指标", "数值"],
        [
            ["比赛总数", "15,817 场"],
            ["训练集（2010-2024）", "14,504 场"],
            ["校准集（2025）", "1,002 场"],
            ["验证集（2026.1.1-2026.6.10）", "311 场"],
            ["涉及球队数", "约 230 支"],
            ["字段数", "8 个"],
            ["时间跨度", "2010-01-01 至 2026-06-10"],
        ])

    add_body(doc, "2026 年世界杯采用 48 队新赛制，分为 12 个小组，每组 4 队。小组前两名和 8 支最佳第三名共 32 队晋级淘汰赛，之后经过 32 强、16 强、8 强、半决赛和决赛决出冠军。分组情况如表 2-2 所示。")

    add_caption(doc, "表 2-2：2026 年世界杯分组")
    groups_data = {}
    groups_csv = os.path.join(PROJECT_ROOT, "data", "worldcup_2026_groups.csv")
    with open(groups_csv, encoding="utf-8") as f:
        for row in csv.DictReader(f):
            groups_data.setdefault(row["group"], []).append(row["team"])

    group_rows = []
    for g in "ABCDEFGHIJKL":
        teams = groups_data.get(g, [])
        group_rows.append([f"{g} 组", "、".join(teams)])
    add_table(doc, ["小组", "球队"], group_rows)

    add_body(doc, "历史比赛数据按年度分布如图 2-1 所示，各年度比赛数量在 850-1200 场之间，2020 年受疫情影响显著减少。")
    add_figure(doc, "fig2_1_matches_per_year.png", "图 2-1：各年度比赛数量趋势（折线图）")

    add_body(doc, "从赛事类型分布来看（图 2-2），友谊赛占比最大，其次是预选赛和洲际杯赛，世界杯正赛占比最小但权重最高。")
    add_figure(doc, "fig2_2_tournament_types.png", "图 2-2：历史比赛赛事类型分布（饼图）", width=Inches(4))

    doc.add_page_break()

    # ===================== CHAPTER 3 =====================
    add_heading(doc, "第三章　探索性分析", level=1)

    add_heading(doc, "3.1 进球数分布", level=2)
    add_body(doc, "对 2010-2024 年历史比赛数据的进球数进行统计分析。单队单场进球均值约 1.35 个，每场比赛总进球均值约 2.7 个，均呈右偏分布，符合泊松分布的基本假设。这一分布特征支持使用泊松回归模型来预测单场比赛进球数。进球数分布如图 3-1 所示。")
    add_figure(doc, "fig3_1_goal_distribution.png", "图 3-1：历史比赛进球数分布（直方图）", width=Inches(5.8))

    add_heading(doc, "3.2 Elo 分差与胜率", level=2)
    add_body(doc, "Elo 评分是衡量球队相对实力的常用指标。本文通过历史比赛记录，采用动态更新方式计算各队 Elo 评分。需要说明的是，历史数据中的 home_team / away_team 只是数据集中的位置标记；世界杯所有比赛均在中立场地进行（东道主除外），模型在模拟世界杯时将 neutral 设为 TRUE，主场优势系数归零。分析发现，Elo 分差与胜率之间存在明显的 S 形正相关关系：Elo 高出 200 分时，胜率约为 75%；高出 400 分时，胜率约为 90%。如图 3-2 所示。")
    add_figure(doc, "fig3_2_elo_vs_winrate.png", "图 3-2：Elo 分差与主队胜率关系（折线图）")

    add_heading(doc, "3.3 变量相关性", level=2)
    add_body(doc, "对模型所用特征变量进行 Pearson 相关分析，结果如图 3-3 所示。Elo 分差与比赛结果的相关性最强（约 0.45），其次是近期胜率差（约 0.20）和近期净胜球差（约 0.18），主场优势也存在一定相关性（约 0.10）。各特征之间不存在严重的多重共线性问题。")
    add_figure(doc, "fig3_3_correlation_heatmap.png", "图 3-3：特征变量相关性热力图", width=Inches(4.5))

    doc.add_page_break()

    # ===================== CHAPTER 4 =====================
    add_heading(doc, "第四章　研究方法", level=1)

    add_heading(doc, "4.1 泊松回归模型（主模型）", level=2)
    add_body(doc, "泊松回归是本文的赛事模拟主模型。足球比赛进球数是非负整数，适合用泊松分布描述。单场进球数 k 的概率为：")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("P(k) = λᵏ · e⁻λ / k!")
    run.font.size = Pt(12)
    run.font.italic = True

    add_body(doc, "其中，λ 表示球队的预期进球数。泊松回归采用 log-link 形式：")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("log(λ) = β₀ + β₁·Elo分差 + β₂·近期进球 + β₃·近期失球 + ...")
    run.font.size = Pt(12)
    run.font.italic = True

    add_body(doc, "模型分别估计双方 90 分钟预期进球数，并通过随机抽样生成比分。模型的输入特征包括：Elo 分差、近 10 场胜率差、主队进攻 vs 客队防守、客队进攻 vs 主队防守、近期净胜球差、主场优势、东道主优势和赛事重要性权重。")

    add_heading(doc, "4.2 二元逻辑回归对照模型", level=2)
    add_body(doc, "逻辑回归预测第一列球队在 90 分钟内获胜与否，用于变量解释、Wald 检验和 ROC/AUC 对照。该模型不生成比分，也不进入世界杯赛事模拟。")

    add_heading(doc, "4.3 蒙特卡洛模拟", level=2)
    add_body(doc, f"蒙特卡洛模拟用于估计赛事路径概率。本次运行将完整赛程模拟 {simulation_count:,} 次，每次从当前已赛状态出发，按真实赛制运行至决赛。剩余小组赛使用泊松模型抽样比分，小组排名按积分、净胜球、进球数排序；淘汰赛 90 分钟平局时，以双方预期进球强度为先验概率决定晋级方。")
    add_body(doc, f"某队夺冠概率 = 该队模拟夺冠次数 / {simulation_count:,}。置信区间根据本次模拟频率的二项分布标准误计算，用于展示蒙特卡洛误差。")

    add_heading(doc, "4.4 滚动窗口特征构造", level=2)
    add_body(doc, "本文采用滚动窗口构造近期状态变量。对每一场历史比赛，只使用该场比赛发生之前的信息，计算双方近 10 场的胜率、平均进球、平均失球和净胜球。该方法可以避免使用未来数据预测过去，从而防止数据泄漏。")

    doc.add_page_break()

    # ===================== CHAPTER 5 =====================
    add_heading(doc, "第五章　模型训练与评价", level=1)

    add_heading(doc, "5.1 模型训练", level=2)
    add_body(doc, "模型训练使用 2010 年 1 月 1 日至 2024 年 12 月 31 日的历史比赛数据；2025 年比赛用于拟合平局概率校准系数；2026 年 1 月 1 日至 6 月 10 日作为独立验证集。泊松回归模型使用 R 语言 glm(..., family = poisson()) 分别拟合双方进球数。")

    add_heading(doc, "5.2 模型评价指标", level=2)
    metrics = read_csv_data("model_metrics.csv")
    metrics_dict = {r["metric"]: r["value"] for r in metrics}

    add_body(doc, f"模型在测试集（{int(float(metrics_dict['samples']))} 场比赛）上的评价指标如表 5-1 所示。")

    add_caption(doc, "表 5-1：模型评价指标")
    add_table(doc,
        ["指标", "数值", "说明"],
        [
            ["验证集样本数", f"{int(float(metrics_dict['samples']))}", "2026.1.1-2026.6.10"],
            ["准确率", f"{float(metrics_dict['accuracy']):.1%}", "胜平负预测正确率"],
            ["Brier Score", f"{float(metrics_dict['brier_score']):.4f}", "越小越好"],
            ["Log Loss", f"{float(metrics_dict['log_loss']):.4f}", "越小越好"],
            ["实际平局率", f"{float(metrics_dict['actual_draw_rate']):.1%}", "验证集观测值"],
            ["平均平局概率", f"{float(metrics_dict['mean_predicted_draw_probability']):.1%}", "校准后预测均值"],
        ])

    add_body(doc, f"模型在独立验证集上的准确率为 {float(metrics_dict['accuracy']):.1%}。准确率只评价最高概率类别；概率质量同时使用 Brier Score、Log Loss 以及实际平局率与平均预测平局概率的差异评价。")

    add_body(doc, "混淆矩阵如表 5-2 所示。")
    add_caption(doc, "表 5-2：混淆矩阵")
    conf = {}
    for r in metrics:
        if r["metric"].startswith("confusion_"):
            conf[r["metric"]] = int(float(r["value"]))
    add_table(doc,
        ["实际 \\ 预测", "主胜(H)", "平局(D)", "客胜(A)"],
        [
            ["主胜(H)", str(conf.get("confusion_H_pred_H", 0)), str(conf.get("confusion_H_pred_D", 0)), str(conf.get("confusion_H_pred_A", 0))],
            ["平局(D)", str(conf.get("confusion_D_pred_H", 0)), str(conf.get("confusion_D_pred_D", 0)), str(conf.get("confusion_D_pred_A", 0))],
            ["客胜(A)", str(conf.get("confusion_A_pred_H", 0)), str(conf.get("confusion_A_pred_D", 0)), str(conf.get("confusion_A_pred_A", 0))],
        ])
    add_body(doc, "最高概率分类很少选择平局，但这不代表模型赋予平局的概率为零。验证集实际平局率与平均预测平局概率应结合 Brier Score 和 Log Loss 一并解释；混淆矩阵仅作为分类阈值下的补充结果。")
    add_figure(doc, "fig5_2_confusion_matrix.png", "图 5-1：混淆矩阵热力图", width=Inches(3.8))

    add_heading(doc, "5.3 模型系数分析", level=2)
    coefs = read_csv_data("model_coefficients.csv")
    add_body(doc, "泊松回归模型的系数如表 5-3 所示。正系数表示该特征增加预期进球数，负系数表示减少。")

    add_caption(doc, "表 5-3：泊松回归模型系数")
    feature_names = {
        "(Intercept)": "截距",
        "elo_diff": "Elo 分差",
        "recent_win_rate_diff": "近期胜率差",
        "home_attack_vs_away_defense": "主队进攻vs客队防守",
        "away_attack_vs_home_defense": "客队进攻vs主队防守",
        "recent_goal_diff_diff": "近期净胜球差",
        "home_field": "主场优势",
        "host_country_diff": "东道主优势",
        "tournament_importance": "赛事重要性",
    }
    coef_rows = []
    for r in coefs:
        fname = feature_names.get(r["feature"], r["feature"])
        coef_rows.append([fname, f"{float(r['home_goal_coefficient']):.4f}", f"{float(r['away_goal_coefficient']):.4f}"])
    add_table(doc, ["特征", "主队进球系数", "客队进球系数"], coef_rows)

    add_body(doc, "从系数可以看出：（1）Elo 分差是最重要的特征，Elo 越高的球队预期进球越多、失球越少；（2）主场优势对主队进球有正向影响、对客队进球有负向影响；（3）东道主优势有一定正向影响，反映了美、加、墨三国作为东道主的小幅优势。模型系数的直观对比如图 5-2 所示。")
    add_figure(doc, "fig5_1_coefficients.png", "图 5-2：泊松回归模型系数对比（柱状图）")

    doc.add_page_break()

    # ===================== CHAPTER 6 =====================
    add_heading(doc, "第六章　预测结果分析", level=1)
    add_body(doc, f"以下结果均来自代码真实运行输出，本次共进行 {simulation_count:,} 次蒙特卡洛模拟，并纳入截至 2026 年 6 月 20 日 09:18 的 30 场已赛小组赛比分。正式报告生成前，程序要求所有关键数据通过确认闸门。")

    add_heading(doc, "6.1 冠军概率预测", level=2)
    champion = read_csv_data("champion_probabilities.csv")
    add_body(doc, "冠军概率 Top 10 如表 6-1 所示。")
    add_caption(doc, "表 6-1：冠军概率 Top 10（含 95% 置信区间）")
    champ_rows = []
    for r in champion:
        prob = float(r["probability"])
        ci_lo = float(r["ci_lower"])
        ci_hi = float(r["ci_upper"])
        champ_rows.append([r["rank"], r["team"], f"{prob:.2%}", f"[{ci_lo:.2%}, {ci_hi:.2%}]"])
    add_table(doc, ["排名", "球队", "夺冠概率", "95% 置信区间"], champ_rows)

    add_figure(doc, "fig6_1_champion_bar.png", "图 6-1：冠军概率 Top 10 柱状图（含95%置信区间误差棒）")
    top3 = champion[:3]
    add_body(
        doc,
        "模型输出的夺冠概率前三名分别为："
        + "、".join(f"{r['team']}（{float(r['probability']):.2%}）" for r in top3)
        + "。该段结论由当前 CSV 输出自动生成，不写死球队名称；若数据或模型重跑，报告应同步更新。"
    )

    add_heading(doc, "6.2 四强概率预测", level=2)
    four = read_csv_data("final_four_probabilities.csv")
    add_body(doc, "四强概率 Top 10 如表 6-2 所示。")
    add_caption(doc, "表 6-2：四强概率 Top 10")
    four_rows = []
    for r in four:
        prob = float(r["probability"])
        four_rows.append([r["rank"], r["team"], f"{prob:.2%}"])
    add_table(doc, ["排名", "球队", "四强概率"], four_rows)
    add_figure(doc, "fig6_2_final_four_bar.png", "图 6-2：四强概率 Top 10 水平柱状图")

    add_heading(doc, "6.3 各小组出线概率", level=2)
    group_adv = read_csv_data("group_advancement_probabilities.csv")
    add_body(doc, "各小组球队的出线概率（晋级 32 强）如表 6-3 所示。")
    add_caption(doc, "表 6-3：各小组出线概率")

    # Group by group
    by_group = {}
    for r in group_adv:
        by_group.setdefault(r["group"], []).append(r)

    ga_rows = []
    for g in "ABCDEFGHIJKL":
        teams = by_group.get(g, [])
        teams.sort(key=lambda x: -float(x["round_of_32"]))
        for t in teams:
            ga_rows.append([
                f"{g} 组", t["team"],
                f"{float(t['round_of_32']):.1%}",
                f"{float(t['round_of_16']):.1%}",
                f"{float(t['champion']):.2%}",
            ])
    add_table(doc, ["小组", "球队", "出线概率", "16强概率", "夺冠概率"], ga_rows)

    add_body(doc, "热门球队从 32 强到夺冠各阶段晋级概率的变化趋势如图 6-3（分组柱状图）和图 6-4（阶梯折线图）所示。可以直观看到，随着赛程推进，各队概率逐轮递减，但强队递减速度明显慢于弱队。")
    add_figure(doc, "fig6_3_progression_funnel.png", "图 6-3：Top 5 球队各阶段晋级概率（分组柱状图）")
    add_figure(doc, "fig6_4_team_funnel_line.png", "图 6-4：热门球队晋级概率阶梯图（折线图）")

    doc.add_page_break()

    add_heading(doc, "6.4 截至 2026-06-20 09:18 的赛中结果对照", level=2)
    add_body(doc, "截至 2026 年 6 月 20 日 09:18，若已赛比分已经完成来源确认，则 30 场小组赛结果如表 6-4 所示。")
    add_caption(doc, "表 6-4：截至 2026-06-20 09:18 已赛结果")
    results_csv = os.path.join(PROJECT_ROOT, "data", "worldcup_2026_results_asof_2026-06-20.csv")
    results = []
    with open(results_csv, encoding="utf-8") as f:
        for r in csv.DictReader(f):
            results.append(r)
    res_rows = []
    for r in results:
        res_rows.append([r["match_date"], f"{r['group']} 组", f"{r['team_1']} vs {r['team_2']}", f"{r['team_1_score']}-{r['team_2_score']}"])
    add_table(doc, ["日期", "小组", "比赛", "比分"], res_rows)

    add_body(doc, "上述赛果通过两条路径影响后续预测：已完成比赛直接进入小组积分与净胜球；同时按比赛顺序更新参赛队 Elo 和近 10 场状态。具体概率变化由 6 月 17 日与 6 月 19 日同随机种子快照对比文件给出，避免凭主观描述判断涨跌。")

    doc.add_page_break()

    # ===================== CHAPTER 7 =====================
    add_heading(doc, "第七章　结论与不足", level=1)

    add_heading(doc, "7.1 研究结论", level=2)
    add_body(doc, f"本文基于公开历史比赛数据、球队 Elo 强度指标和经确认的世界杯赛程，构建泊松回归单场进球预测模型，并通过 {simulation_count:,} 次蒙特卡洛模拟估计各队晋级概率和冠军概率。主要结论如下：")
    add_body(
        doc,
        "（1）模型输出的夺冠概率前三名为"
        + "、".join(f"{r['team']}（{float(r['probability']):.2%}）" for r in top3)
        + "。该结论依赖已经确认的数据源和本次代码输出。"
    )
    add_body(doc, "（2）Elo 分差是预测比赛结果最重要的单一特征，其模型系数显著高于其他特征。")
    add_body(doc, "（3）48 队新赛制下，由于增加了最佳第三名晋级机制，传统强队的出线概率普遍较高（>85%），但淘汰赛路径的随机性显著影响最终排名。")
    add_body(doc, f"（4）模型在测试集上的准确率为 {float(metrics_dict['accuracy']):.1%}，对于足球比赛胜平负三分类预测属于合理水平。")

    add_heading(doc, "7.2 研究不足与改进方向", level=2)
    add_body(doc, "（1）本文训练数据以公开历史比赛结果为核心，未包含球员级别信息、临场伤病、战术变化、天气条件和市场赔率等因素，模型解释力受限。")
    add_body(doc, "（2）泊松回归假设双方进球独立，但实际比赛中低比分情况可能存在相关性。未来可引入 Dixon-Coles 修正改进低比分预测。")
    add_body(doc, "（3）2026 年世界杯由三国共同举办，传统主场优势模型可能不完全适用。未来可加入地理距离、时区差异等变量。")
    add_body(doc, "（4）最高概率分类仍较少选择平局；本文已使用独立的 2025 年校准集调整平局概率，但仍可进一步采用 Dixon-Coles 联合比分模型改善低比分相关性。")

    doc.add_page_break()

    # ===================== REFERENCES =====================
    add_heading(doc, "参考文献", level=1)
    refs = [
        "[1] FIFA. FIFA World Cup 2026 official tournament page. https://www.fifa.com/en/tournaments/mens/worldcup/canadamexicousa2026",
        "[2] FIFA. Regulations for the FIFA World Cup 2026. FIFA official regulations document.",
        "[3] Maher, M. J. Modelling association football scores. Statistica Neerlandica, 1982.",
        "[4] Dixon, M. J., Coles, S. G. Modelling association football scores and inefficiencies in the football betting market. Journal of the Royal Statistical Society: Series C, 1997.",
        "[5] World Football Elo Ratings. https://www.eloratings.net/",
        "[6] International Football Results from 1872 to Present. https://github.com/martj42/international_results",
        "[7] 百度体育. 世界杯赛程与排名页. https://tiyu.baidu.com",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(ref)
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    # ===================== APPENDIX =====================
    add_heading(doc, "附录 A：核心代码文件说明", level=1)

    add_caption(doc, "表 A-1：核心代码文件")
    add_table(doc,
        ["文件", "作用"],
        [
            ["R/validate_inputs.R", "校验分组、赛程、历史训练集字段和日期边界"],
            ["R/worldcup_predictor.R", "完整建模与模拟主程序（约 600 行）"],
            ["R/make_charts.R", "根据模型输出结果生成 SVG 图表"],
        ])

    add_body(doc, "运行命令：")
    p = doc.add_paragraph()
    run = p.add_run("Rscript R/validate_inputs.R\nRscript R/worldcup_predictor.R\nRscript R/make_charts.R")
    run.font.size = Pt(9)
    run.font.name = "Courier New"

    doc.add_page_break()

    # ---------- Appendix B: Full R code ----------
    add_heading(doc, "附录 B：worldcup_predictor.R 完整代码", level=1)
    r_code_path = os.path.join(PROJECT_ROOT, "R", "worldcup_predictor.R")
    with open(r_code_path, encoding="utf-8") as f:
        r_code = f.read()
    add_code_block(doc, r_code)

    doc.add_page_break()

    add_heading(doc, "附录 C：validate_inputs.R 完整代码", level=1)
    val_code_path = os.path.join(PROJECT_ROOT, "R", "validate_inputs.R")
    with open(val_code_path, encoding="utf-8") as f:
        val_code = f.read()
    add_code_block(doc, val_code)

    doc.add_page_break()

    add_heading(doc, "附录 D：make_charts.R 完整代码", level=1)
    chart_code_path = os.path.join(PROJECT_ROOT, "R", "make_charts.R")
    if os.path.exists(chart_code_path):
        with open(chart_code_path, encoding="utf-8") as f:
            chart_code = f.read()
        add_code_block(doc, chart_code)

    # Save
    output_path = os.path.join(PROJECT_ROOT, "2026世界杯冠军概率预测报告.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    main()
