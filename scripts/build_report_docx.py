#!/usr/bin/env python3
"""Build a DOCX version of the 2026 World Cup report draft."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / "2026世界杯冠军概率预测报告_草稿.md"
OUTPUT = ROOT / "2026世界杯冠军概率预测报告.docx"
CJK_FONT = "PingFang SC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", stripped):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    widths = [2400] + [int((9360 - 2400) / max(1, max_cols - 1))] * (max_cols - 1)
    for r_idx, row in enumerate(rows):
        set_row_cant_split(table.rows[r_idx])
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[min(c_idx, len(widths) - 1)])
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
                    run.font.size = Pt(9.5)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()


def add_inline_text(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)
        run.font.name = run.font.name or "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build() -> None:
    doc = Document()
    configure_styles(doc)

    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            flush_table()
            continue
        if line.startswith("|"):
            table_buffer.append(line)
            continue

        flush_table()
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor.from_string("0B2545")
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            add_inline_text(p, line)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_text(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_inline_text(p, line)

    flush_table()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("2026 年世界杯冠军概率预测研究")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor.from_string("666666")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
